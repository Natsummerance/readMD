# -*- coding: utf-8 -*-
import io, json, os, sys, tempfile, threading, unittest, urllib.request
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

import readmd_modules.mdcheck as MDC
from readmd_modules import convert as CV


class MockAI(BaseHTTPRequestHandler):
    """本地 mock：四种协议 + /models。"""
    def log_message(self, *a):
        pass

    def _json(self, obj, code=200):
        body = json.dumps(obj, ensure_ascii=False).encode('utf-8')
        self.send_response(code)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        if self.path in ('/v1/models', '/models'):
            self._json({'data': [{'id': 'mock-a'}, {'id': 'mock-b'}]})
        else:
            self._json({'error': 'nf'}, 404)

    def do_POST(self):
        n = int(self.headers.get('Content-Length', 0))
        body = json.loads(self.rfile.read(n) or b'{}')
        path = self.path
        stream = body.get('stream', False)

        def ev(chunks):
            self.send_response(200)
            self.send_header('Content-Type', 'text/event-stream')
            self.end_headers()
            for c in chunks:
                self.wfile.write(('data: ' + json.dumps(c) + '\n\n').encode())
            self.wfile.write(b'data: [DONE]\n\n')

        if path.endswith('/chat/completions'):
            if stream:
                ev([{'choices': [{'delta': {'content': 'Hi '}}]},
                    {'choices': [{'delta': {'content': 'there'}}]},
                    {'usage': {'prompt_tokens': 10, 'completion_tokens': 5, 'total_tokens': 15}},
                    {'choices': [{'delta': {}, 'finish_reason': 'stop'}]}])
            else:
                self._json({'choices': [{'message': {'content': 'Hi'}}],
                            'usage': {'prompt_tokens': 8, 'completion_tokens': 4, 'total_tokens': 12}})
        elif path.endswith('/completions'):
            self._json({'choices': [{'text': 'Comp answer'}],
                        'usage': {'prompt_tokens': 3, 'completion_tokens': 2, 'total_tokens': 5}})
        elif path.endswith('/responses'):
            if stream:
                ev([{'type': 'response.output_text.delta', 'delta': {'text': 'Resp '}},
                    {'type': 'response.output_text.delta', 'delta': {'text': 'ok'}},
                    {'type': 'response.completed', 'response': {'usage': {'prompt_tokens': 7, 'completion_tokens': 3, 'total_tokens': 10}}}])
            else:
                self._json({'output_text': 'Resp ok',
                            'usage': {'prompt_tokens': 7, 'completion_tokens': 3, 'total_tokens': 10}})
        elif path.endswith('/v1/messages'):
            if stream:
                ev([{'type': 'message_start', 'message': {'usage': {'input_tokens': 11}}},
                    {'type': 'content_block_delta', 'delta': {'type': 'text_delta', 'text': 'Anth '}},
                    {'type': 'content_block_delta', 'delta': {'type': 'text_delta', 'text': 'ok'}},
                    {'type': 'message_delta', 'usage': {'output_tokens': 6}},
                    {'type': 'message_stop'}])
            else:
                self._json({'content': [{'type': 'text', 'text': 'Anth ok'}],
                            'usage': {'input_tokens': 11, 'output_tokens': 6}})
        else:
            self._json({'error': 'nf'}, 404)


class TestDocx(unittest.TestCase):
    def test_docx2md_rich(self):
        from docx import Document
        from docx.oxml.ns import qn
        import lxml.etree as ET
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'sample.docx')
            d = Document()
            d.add_heading('测试标题', level=1)
            para = d.add_paragraph()
            omml = (
                '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                '<m:oMath><m:f><m:num><m:r><m:t>a</m:t></m:r></m:num>'
                '<m:den><m:r><m:t>b</m:t></m:r></m:den></m:f></m:oMath></m:oMathPara>'
            )
            para._p.append(ET.fromstring(omml))
            t = d.add_table(rows=2, cols=2)
            t.cell(0, 0).text = '列A'
            t.cell(0, 1).text = '列B'
            t.cell(1, 0).text = '1'
            t.cell(1, 1).text = '2'
            cp = d.add_paragraph()
            run = cp.add_run('print("hi")')
            run.font.name = 'Consolas'
            d.add_paragraph('项目一', style='List Bullet')
            d.save(p)

            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('# 测试标题', text)
            self.assertIn(r'\frac{a}{b}', text)
            self.assertIn('| 列A | 列B |', text)
            self.assertIn('```', text)
            self.assertIn('print("hi")', text)
            self.assertIn('- 项目一', text)


class TestPdf(unittest.TestCase):
    def test_pdf2md_table(self):
        import fitz
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 't.pdf')
            doc = fitz.open()
            page = doc.new_page()
            x0, y0, x1, y1 = 72, 72, 280, 120
            rows, cols = 2, 2
            cw = (x1 - x0) / cols
            ch = (y1 - y0) / rows
            for r in range(rows + 1):
                page.draw_line((x0, y0 + r * ch), (x1, y0 + r * ch), color=(0, 0, 0), width=0.6)
            for c in range(cols + 1):
                page.draw_line((x0 + c * cw, y0), (x0 + c * cw, y1), color=(0, 0, 0), width=0.6)
            page.insert_text((x0 + 4, y0 + 14), 'H1', fontsize=9)
            page.insert_text((x0 + cw + 4, y0 + 14), 'H2', fontsize=9)
            page.insert_text((x0 + 4, y0 + ch + 14), 'v1', fontsize=9)
            page.insert_text((x0 + cw + 4, y0 + ch + 14), 'v2', fontsize=9)
            page.insert_text((72, 150), 'Text after table line', fontsize=11)
            doc.save(p)
            doc.close()

            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'pdf', err)
            self.assertIn('|', text)
            self.assertIn('H1', text)
            self.assertIn('H2', text)
            self.assertIn('Text after table line', text)


class TestMdcheck(unittest.TestCase):
    def test_fence_close(self):
        fixed, issues = MDC.check('```python\nx=1')
        self.assertIn('```', fixed)
        self.assertTrue(any('围栏' in i['msg'] for i in issues))

    def test_math_delim(self):
        fixed, issues = MDC.check('$$x$$ 和 $y')
        self.assertTrue(any('公式' in i['msg'] and '奇数' in i['msg'] for i in issues))

    def test_replace_char(self):
        fixed, issues = MDC.check('a\uFFFDb')
        self.assertTrue(any('替换符' in i['msg'] for i in issues))

    def test_missing_image(self):
        with tempfile.TemporaryDirectory() as td:
            fixed, issues = MDC.check('![x](missing.png)', td)
            self.assertTrue(any('图片引用不存在' in i['msg'] for i in issues))

    def test_table_fix(self):
        fixed, issues = MDC.check('| A | B |\n| 1 | 2 |')
        self.assertIn('| --- | --- |', fixed)

    def test_blank_fold(self):
        fixed, issues = MDC.check('a\n\n\n\n\nb')
        self.assertEqual(fixed.count('\n\n'), 1)


class TestConvertApi(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        from readmd import Handler, RM
        RM.load_forced('convert')
        cls.srv = ThreadingHTTPServer(('127.0.0.1', 0), Handler)
        cls.srv.daemon_threads = True
        threading.Thread(target=cls.srv.serve_forever, daemon=True).start()
        cls.base = 'http://127.0.0.1:%d' % cls.srv.server_port

    @classmethod
    def tearDownClass(cls):
        cls.srv.shutdown()
        cls.srv.server_close()

    def _mk_docx(self, td, name='a'):
        from docx import Document
        p = os.path.join(td, name + '.docx')
        d = Document()
        d.add_heading('标题%s' % name, level=1)
        d.add_paragraph('正文%s' % name)
        d.save(p)
        return p

    def test_single_autosave(self):
        with tempfile.TemporaryDirectory() as td:
            p = self._mk_docx(td)
            with urllib.request.urlopen(self.base + '/api/convert?p=' + urllib.request.quote(p)) as r:
                d = json.loads(r.read().decode('utf-8'))
            self.assertTrue(d['saved'], d)
            self.assertEqual(d['engine'], 'docx')
            self.assertTrue(os.path.isfile(d['out']))
            self.assertIn('标题a', open(d['out'], encoding='utf-8').read())

    def test_single_skip_when_exists(self):
        with tempfile.TemporaryDirectory() as td:
            p = self._mk_docx(td, 'b')
            out = os.path.join(td, 'b.md')
            with open(out, 'w', encoding='utf-8') as f:
                f.write('existing')
            with urllib.request.urlopen(self.base + '/api/convert?p=' + urllib.request.quote(p)) as r:
                d = json.loads(r.read().decode('utf-8'))
            self.assertTrue(d['skipped'], d)
            self.assertEqual(open(out, encoding='utf-8').read(), 'existing')

    def test_batch_and_overwrite(self):
        import time
        with tempfile.TemporaryDirectory() as td:
            p1 = self._mk_docx(td, 'x1')
            p2 = self._mk_docx(td, 'x2')
            req = urllib.request.Request(
                self.base + '/api/convert/batch',
                data=json.dumps({'paths': [p1, p2], 'overwrite': False}).encode('utf-8'),
                headers={'Content-Type': 'application/json'}, method='POST')
            with urllib.request.urlopen(req, timeout=30) as r:
                bd = json.loads(r.read().decode('utf-8'))
            jid = bd['job']
            pr = self._poll(jid)
            self.assertTrue(pr['finished'], pr)
            self.assertTrue(all(i['status'] == 'ok' for i in pr['items']), pr)
            self.assertTrue(os.path.isfile(os.path.join(td, 'x1.md')))
            pr2 = self._poll(self._start([p1], False))
            self.assertEqual(pr2['items'][0]['status'], 'skipped', pr2)
            pr3 = self._poll(self._start([p1], True))
            self.assertEqual(pr3['items'][0]['status'], 'ok', pr3)

    def _start(self, paths, overwrite):
        req = urllib.request.Request(
            self.base + '/api/convert/batch',
            data=json.dumps({'paths': paths, 'overwrite': overwrite}).encode('utf-8'),
            headers={'Content-Type': 'application/json'}, method='POST')
        with urllib.request.urlopen(req, timeout=30) as r:
            return json.loads(r.read().decode('utf-8'))['job']

    def _poll(self, jid):
        import time
        pr = {}
        for _ in range(80):
            with urllib.request.urlopen(self.base + '/api/convert/progress?job=' + jid) as r:
                pr = json.loads(r.read().decode('utf-8'))
            if pr.get('finished'):
                break
            time.sleep(0.25)
        return pr

    def test_collect(self):
        with tempfile.TemporaryDirectory() as td:
            for n in ('a.docx', 'b.pdf', 'c.md', 'd.txt'):
                with open(os.path.join(td, n), 'w', encoding='utf-8') as f:
                    f.write('x')
            with urllib.request.urlopen(self.base + '/api/convert/collect?dir=' + urllib.request.quote(td)) as r:
                d = json.loads(r.read().decode('utf-8'))
            names = sorted(os.path.basename(x) for x in d['files'])
            self.assertEqual(names, ['a.docx', 'b.pdf', 'd.txt'])


class TestAi(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        from readmd_modules import ai
        cls.ai = ai
        cls.mock = ThreadingHTTPServer(('127.0.0.1', 0), MockAI)
        threading.Thread(target=cls.mock.serve_forever, daemon=True).start()
        cls.base = 'http://127.0.0.1:%d' % cls.mock.server_port

    @classmethod
    def tearDownClass(cls):
        cls.mock.shutdown()
        cls.mock.server_close()

    def _run(self, payload):
        return list(self.ai.chat(payload))

    def test_chat_stream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'chat', 'stream': True})
        self.assertEqual(''.join(x for x in r if isinstance(x, str)), 'Hi there')
        u = [x for x in r if isinstance(x, dict)][0]['usage']
        self.assertEqual(u, {'prompt_tokens': 10, 'completion_tokens': 5, 'total_tokens': 15})

    def test_chat_nostream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'chat', 'stream': False})
        self.assertEqual(r[0], 'Hi')
        self.assertEqual(r[1]['usage']['total_tokens'], 12)

    def test_completion(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'completion', 'stream': False})
        self.assertEqual(r[0], 'Comp answer')
        self.assertEqual(r[1]['usage']['total_tokens'], 5)

    def test_responses_stream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'responses', 'stream': True})
        self.assertEqual(''.join(x for x in r if isinstance(x, str)), 'Resp ok')
        self.assertEqual([x for x in r if isinstance(x, dict)][0]['usage']['total_tokens'], 10)

    def test_responses_nostream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'responses', 'stream': False})
        self.assertEqual(r[0], 'Resp ok')
        self.assertEqual(r[1]['usage']['total_tokens'], 10)

    def test_anthropic_stream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'messages', 'stream': True})
        self.assertEqual(''.join(x for x in r if isinstance(x, str)), 'Anth ok')
        u = [x for x in r if isinstance(x, dict)][0]['usage']
        self.assertEqual(u, {'prompt_tokens': 11, 'completion_tokens': 6, 'total_tokens': 17})

    def test_anthropic_nostream(self):
        r = self._run({'base_url': self.base, 'api_key': 'k', 'model': 'm',
                       'messages': [{'role': 'user', 'content': 'hi'}],
                       'mode': 'messages', 'stream': False})
        self.assertEqual(r[0], 'Anth ok')
        self.assertEqual(r[1]['usage']['total_tokens'], 17)

    def test_list_models(self):
        ids = self.ai.list_models(self.base, 'k', 'auto')
        self.assertEqual(ids, ['mock-a', 'mock-b'])
        ids2 = self.ai.list_models(self.base, 'k', 'messages')
        self.assertEqual(ids2, ['mock-a', 'mock-b'])
        with self.assertRaises(self.ai.ChatError):
            self.ai.list_models('', '', 'auto')

    def test_get_config_mode_field(self):
        cfg = self.ai.get_config()
        self.assertIn('presets', cfg)
        self.assertIn('mode', cfg['presets'][0])
        anth = next((p for p in cfg['presets'] if p.get('format') == 'anthropic'), None)
        if anth:
            self.assertEqual(anth['mode'], 'messages')

    def test_config_v2_resets_legacy_and_masks_key(self):
        """旧格式升级时清空自定义项；配置接口绝不返回保存的 Key。"""
        path = self.ai.CONFIG_FILE
        old = None
        if os.path.isfile(path):
            with open(path, 'rb') as f:
                old = f.read()
        try:
            self.ai._write_cfg({'providers': [{'name': 'old connection', 'api_key': 'secret'}],
                                'current': {'provider': 'old connection', 'model': 'old'}})
            migrated = self.ai.get_config()
            self.assertEqual(migrated['custom'], [])
            self.assertEqual(migrated['current'], {})

            self.ai.save_config({'providers': [{'name': 'my connection', 'base_url': self.base,
                                                'mode': 'auto', 'models': ['mock-a'], 'api_key': 'secret'}],
                                 'current': {'provider': 'my connection', 'model': 'mock-a'}})
            public = self.ai.get_config()
            custom = public['custom'][0]
            self.assertTrue(custom['id'].startswith('custom:'))
            self.assertTrue(custom['has_key'])
            self.assertNotIn('api_key', custom)
            self.assertEqual(self.ai.resolve_key(self.ai.find_provider('my connection')), 'secret')

            # 空 Key 表示“不改动”，避免编辑 URL/模型时意外丢失密钥。
            self.ai.save_config({'providers': [{'name': 'my connection', 'base_url': self.base,
                                                'mode': 'auto', 'models': ['mock-b']}],
                                 'current': {'provider': 'my connection', 'model': 'mock-b'}})
            self.assertEqual(self.ai.resolve_key(self.ai.find_provider('my connection')), 'secret')

            # 重命名依赖稳定 ID，且 clear_key 必须显式删除密钥。
            provider_id = self.ai.get_config()['custom'][0]['id']
            self.ai.save_config({'providers': [{'id': provider_id, 'name': 'renamed connection',
                                                'base_url': self.base, 'mode': 'auto',
                                                'models': ['mock-b'], 'clear_key': True}],
                                 'current': {'provider_id': provider_id, 'model': 'mock-b'}})
            renamed = self.ai.get_config()
            self.assertEqual(renamed['custom'][0]['id'], provider_id)
            self.assertFalse(renamed['custom'][0]['has_key'])
            self.assertEqual(renamed['current']['provider_id'], provider_id)
        finally:
            if old is None:
                try:
                    os.unlink(path)
                except OSError:
                    pass
            else:
                with open(path, 'wb') as f:
                    f.write(old)


def main():
    suite = unittest.defaultTestLoader.loadTestsFromModule(sys.modules[__name__])
    result = unittest.TextTestRunner(verbosity=1).run(suite)
    sys.exit(0 if result.wasSuccessful() else 1)


if __name__ == '__main__':
    main()
