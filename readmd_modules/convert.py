# -*- coding: utf-8 -*-
"""万物转 md：核心格式专用解析（docx / pdf）+ MarkItDown 兜底 + 严格校验。

v2.1.1 质量升级：
- .docx：python-docx + lxml 专用解析 —— OMML 公式转 LaTeX、标题层级、表格、
  等宽字体代码块、图片引用；
- .pdf：PyMuPDF 专用解析 —— find_tables 还原边框表格 + 公式启发式；
- 其余格式（pptx/xlsx/html/csv/json/zip 等）走 MarkItDown；
- 专用解析异常时逐文件回退 MarkItDown，仍失败抛出带原因的异常。
"""

import os
import re

_engine = None

_MONO_FONTS = ('consolas', 'courier new', 'courier', 'menlo', 'monaco',
               'source code pro', 'cascadia code', 'fira code', 'jetbrains mono',
               'sf mono', 'liberation mono', 'dejavu sans mono')

_CODE_LANG_HINTS = (
    ('python', 'python'), ('javascript', 'javascript'), ('typescript', 'typescript'),
    ('java', 'java'), ('csharp', 'csharp'), ('c++', 'cpp'), ('cpp', 'cpp'),
    ('go', 'go'), ('rust', 'rust'), ('sql', 'sql'), ('html', 'html'),
    ('css', 'css'), ('bash', 'bash'), ('shell', 'bash'), ('json', 'json'),
    ('xml', 'xml'), ('yaml', 'yaml'), ('powershell', 'powershell'),
    ('php', 'php'), ('ruby', 'ruby'), ('vb', 'vb'), ('swift', 'swift'),
    ('kotlin', 'kotlin'), ('c#', 'csharp'), ('c', 'c'), ('cs', 'csharp'),
)

_MATH_CHARS = set('+-*/=<>^_~%∑∫√∞≈≠±×÷∂∇∏πθαβγδφψωλμνστρΔΩΦΓΛεηζξοπς')


def load():
    global _engine
    if _engine is None:
        from markitdown import MarkItDown
        _engine = MarkItDown()
    return _engine


def supported_hint():
    return ('支持：PDF / Word / PowerPoint / Excel / HTML / CSV / JSON / XML / '
            '邮件 / 压缩包等；图片与扫描件请用 OCR 模块')


# ---------------------------------------------------------------- 入口

def convert(path):
    """把任意支持的文件转换为 Markdown 文本（专用解析 → MarkItDown 兜底）。"""
    text, _engine_name, error = convert_verbose(path)
    if error and not text:
        raise ValueError(error)
    return text


def convert_verbose(path):
    """返回 (text, engine, error)。engine: 'docx' | 'pdf' | 'markitdown' | ''"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        try:
            return docx2md(path), 'docx', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    if ext == '.pdf':
        try:
            return pdf2md(path), 'pdf', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    try:
        return _markitdown_convert(path), 'markitdown', None
    except Exception as e:  # noqa: BLE001
        return '', '', str(e)


def _markitdown_convert(path):
    eng = load()
    result = eng.convert(path)
    return (result.text_content or '').strip()


# ---------------------------------------------------------------- docx 专用

_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _mq(tag):
    return '{%s}%s' % (_M_NS, tag)


def _omml_to_latex(el):
    """递归把 OMML 元素转为 LaTeX 片段（基础结构：分数/上下标/根号/求和积分等）。"""
    if el is None:
        return ''
    tag = el.tag
    if not isinstance(tag, str) or not tag.startswith('{'):
        return (el.text or '')
    local = tag.split('}')[1]

    def kids(e):
        return ''.join(_omml_to_latex(c) for c in e)

    if local == 't':
        return (el.text or '').replace('\u200b', '').replace('\u2061', '')
    if local in ('r', 'oMath', 'oMathPara', 'm', 'e', 'num', 'den', 'sub',
                 'sup', 'deg', 'chr', 'fName', 'acc', 'd', 'delim', 'eqArr',
                 'limLow', 'limUpp', 'groupChr', 'bar', 'phant', 'box',
                 'borderBox', 'func'):
        return kids(el)
    if local == 'f':  # 分数
        num = el.find(_mq('num'))
        den = el.find(_mq('den'))
        return r'\frac{%s}{%s}' % (kids(num), kids(den))
    if local == 'sSup':
        base = el.find(_mq('e'))
        sup = el.find(_mq('sup'))
        return '{%s}^{%s}' % (kids(base), kids(sup))
    if local == 'sSub':
        base = el.find(_mq('e'))
        sub = el.find(_mq('sub'))
        return '{%s}_{%s}' % (kids(base), kids(sub))
    if local == 'sSubSup':
        base = el.find(_mq('e'))
        sub = el.find(_mq('sub'))
        sup = el.find(_mq('sup'))
        return '{%s}_{%s}^{%s}' % (kids(base), kids(sub), kids(sup))
    if local == 'rad':  # 根式
        deg = el.find(_mq('deg'))
        e_el = el.find(_mq('e'))
        inner = kids(e_el)
        d = kids(deg).strip()
        if d:
            return r'\sqrt[%s]{%s}' % (d, inner)
        return r'\sqrt{%s}' % inner
    if local == 'nary':  # 求和/积分/连乘
        chr_el = el.find(_mq('chr'))
        sub = el.find(_mq('sub'))
        sup = el.find(_mq('sup'))
        e_el = el.find(_mq('e'))
        c = kids(chr_el)
        op = {'\u2211': r'\sum', '\u222b': r'\int', '\u220f': r'\prod',
              '\u222e': r'\oint', '\u22c2': r'\bigcap', '\u22c3': r'\bigcup'}.get(c, c)
        s, p = kids(sub), kids(sup)
        inner = kids(e_el)
        if s or p:
            return '%s_{%s}^{%s} %s' % (op, s, p, inner)
        return '%s %s' % (op, inner)
    if local == 'func':  # 函数
        fname = el.find(_mq('fName'))
        e_el = el.find(_mq('e'))
        return r'%s(%s)' % (kids(fname), kids(e_el))
    if local == 'acc':  # 帽子/箭头
        acc = el.find(_mq('acc'))
        e_el = el.find(_mq('e'))
        a = kids(acc)
        inner = kids(e_el)
        marks = {'\u02c6': r'\hat', '\u00af': r'\bar', '\u2192': r'\vec',
                 '\u02dc': r'\tilde', '\u0307': r'\dot', '\u0308': r'\ddot'}
        cmd = marks.get(a, '')
        return (cmd + '{%s}' % inner) if cmd else inner
    if local == 'bar':  # 上下横线
        pos = el.find(_mq('barPr'))
        e_el = el.find(_mq('e'))
        return r'\overline{%s}' % kids(e_el)
    if local == 'groupChr':  # 括号/花括号
        chr_el = el.find(_mq('chr'))
        e_el = el.find(_mq('e'))
        c = kids(chr_el)
        pairs = {'{': r'\left\{ %s \right\}', '}': r'\left\{ %s \right\}',
                 '[': r'\left[ %s \right]', ']': r'\left[ %s \right]',
                 '(': r'\left( %s \right)', ')': r'\left( %s \right)',
                 '|': r'\left| %s \right|'}
        if c in pairs:
            return pairs[c] % kids(e_el)
        return kids(e_el)
    if local in ('rPr', 'ctrlPr', 'argPr', 'eqArrPr', 'naryPr', 'sSupPr',
                 'sSubPr', 'sSubSupPr', 'radPr', 'fPr', 'accPr', 'barPr',
                 'delimPr', 'funcPr', 'limLowPr', 'limUppPr', 'groupChrPr',
                 'phantPr', 'boxPr', 'borderBoxPr', 'mathPr', 'wrapPr',
                 'intLim', 'naryLim', 'subHide', 'supHide'):
        return ''
    # 未知标签：取文本兜底
    return kids(el)


def _run_font_lower(r):
    try:
        name = (r.font.name or '') or ''
    except Exception:  # noqa: BLE001
        name = ''
    if not name:
        rPr = r._r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            rf = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rf is not None:
                name = (rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                        or rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')
                        or '')
    return (name or '').lower()


def _para_has_mono(p):
    return any(_run_font_lower(r) in _MONO_FONTS for r in p.runs if (r.text or '').strip())


def _para_plain(p):
    return ''.join((r.text or '') for r in p.runs)


def _para_inline(p):
    parts = []
    for r in p.runs:
        t = r.text or ''
        if not t:
            continue
        if _run_font_lower(r) in _MONO_FONTS:
            t = '`' + t + '`'
        if r.bold:
            t = '**' + t + '**'
        if r.italic:
            t = '*' + t + '*'
        parts.append(t)
    return ''.join(parts)


def _lang_hint(text):
    low = (text or '').lower()
    for key, lang in _CODE_LANG_HINTS:
        if key in low:
            return lang
    return ''


def _docx_equations(p):
    """返回段落中所有 m:oMath 元素转出的 LaTeX 片段。"""
    out = []
    for om in p._p.iter(_mq('oMath')):
        latex = _omml_to_latex(om).strip()
        if latex:
            out.append(latex)
    return out


def _table_to_md(tbl):
    """w:tbl → 规整管道表（按实际 tc 遍历，处理合并单元格去重）。"""
    from docx.table import _Cell
    rows = []
    for tr in tbl._tbl.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
        cells = []
        seen = set()
        for tc in tr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
            if id(tc) in seen:
                continue
            seen.add(id(tc))
            c = _Cell(tc, tbl)
            txt = '\n'.join(pp.text or '' for pp in c.paragraphs)
            txt = txt.replace('\n', ' ').replace('|', '\\|').strip()
            cells.append(txt)
        rows.append(cells)
    if not rows:
        return ''
    ncol = max(len(r) for r in rows)
    rows = [r + [''] * (ncol - len(r)) for r in rows]
    out = ['| ' + ' | '.join(rows[0]) + ' |',
           '| ' + ' | '.join(['---'] * ncol) + ' |']
    for r in rows[1:]:
        out.append('| ' + ' | '.join(r) + ' |')
    return '\n'.join(out)


def docx2md(path):
    """docx → Markdown：OMML 公式→LaTeX、标题层级、表格、等宽字体代码块。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    base_dir = os.path.dirname(os.path.abspath(path))
    lines = []
    code_buf = []
    code_lang = ''

    def flush_code():
        nonlocal code_buf, code_lang
        if code_buf:
            lines.append('```' + code_lang)
            lines.extend(code_buf)
            lines.append('```')
            lines.append('')
            code_buf, code_lang = [], ''

    def handle_para(p):
        nonlocal code_buf, code_lang
        eqs = _docx_equations(p)
        if eqs:
            flush_code()
            text = _para_plain(p).strip()
            for latex in eqs:
                if len(latex) >= 80 or _para_plain(p).strip() == latex:
                    lines.append('$$%s$$' % latex)
                else:
                    lines.append('$%s$' % latex)
            if text and text not in eqs:
                lines.append(text)
            lines.append('')
            return
        style_name = ''
        try:
            style_name = (p.style.name or '') if p.style is not None else ''
        except Exception:  # noqa: BLE001
            style_name = ''
        if style_name and style_name.lower().startswith('heading'):
            flush_code()
            try:
                level = int(''.join(ch for ch in style_name if ch.isdigit()) or '1')
            except ValueError:
                level = 1
            level = max(1, min(6, level))
            txt = _para_inline(p).strip()
            if txt:
                lines.append('#' * level + ' ' + txt)
                lines.append('')
            return
        if style_name and style_name.lower() == 'title':
            flush_code()
            txt = _para_inline(p).strip()
            if txt:
                lines.append('# ' + txt)
                lines.append('')
            return
        if _para_has_mono(p):
            txt = _para_plain(p).strip()
            if txt:
                if not code_lang:
                    code_lang = _lang_hint(txt)
                code_buf.append(txt)
            return
        flush_code()
        txt = _para_inline(p).strip()
        if not txt:
            return
        # 列表：段落级或样式级 numPr → "- "
        is_list = False
        ppr = p._p.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:numPr')) is not None:
            is_list = True
        else:
            try:
                st = p.style
                depth = 0
                while st is not None and depth < 8:
                    spPr = st.element.find(qn('w:pPr'))
                    if spPr is not None and spPr.find(qn('w:numPr')) is not None:
                        is_list = True
                        break
                    st = st.base_style
                    depth += 1
            except Exception:  # noqa: BLE001
                is_list = False
        if is_list:
            lines.append('- ' + txt)
            lines.append('')
            return
        lines.append(txt)
        lines.append('')

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            handle_para(Paragraph(child, doc))
        elif child.tag == qn('w:tbl'):
            flush_code()
            md = _table_to_md(Table(child, doc))
            if md:
                lines.append(md)
                lines.append('')
        elif child.tag == qn('w:sectPr'):
            continue
    flush_code()
    text = '\n'.join(lines).strip()
    if not text:
        raise ValueError('docx 未提取到文字内容')
    return text


# ---------------------------------------------------------------- pdf 专用

def _data_to_md(data):
    """表格二维数组 → 管道表。"""
    if not data:
        return ''
    ncol = max(len(r) for r in data)
    data = [r + [''] * (ncol - len(r)) for r in data]

    def clean(cell):
        s = (cell or '') if isinstance(cell, str) else str(cell or '')
        s = s.replace('\n', ' ').replace('\r', ' ').replace('|', '\\|').strip()
        return s

    out = ['| ' + ' | '.join(clean(c) for c in data[0]) + ' |',
           '| ' + ' | '.join(['---'] * ncol) + ' |']
    for r in data[1:]:
        out.append('| ' + ' | '.join(clean(c) for c in r) + ' |')
    return '\n'.join(out)


def _looks_like_formula(line):
    s = line.strip()
    if not s or len(s) < 2 or len(s) > 160:
        return False
    if any('\u4e00' <= ch <= '\u9fff' for ch in s):
        return False
    sig = sum(1 for ch in s if ch in _MATH_CHARS or ch.isdigit() or ch in '()[]{},.')
    return sig / max(len(s), 1) >= 0.45


def _page_to_md(page):
    try:
        import fitz
    except Exception:  # noqa: BLE001
        return ''
    tables = []
    try:
        tables = list(page.find_tables().tables)
    except Exception:  # noqa: BLE001
        tables = []
    tbl_boxes = []
    for t in tables:
        try:
            tbl_boxes.append(fitz.Rect(t.bbox))
        except Exception:  # noqa: BLE001
            pass

    def in_table(r):
        try:
            area = r.get_area()
            if area <= 0:
                return False
            for tb in tbl_boxes:
                inter = r & tb
                if not inter.is_empty and inter.get_area() > 0.5 * area:
                    return True
        except Exception:  # noqa: BLE001
            return False
        return False

    items = []
    seq = 0
    for t in tables:
        try:
            md = _data_to_md(t.extract())
        except Exception:  # noqa: BLE001
            md = ''
        if md:
            items.append((t.bbox[1], seq, 'table', md))
            seq += 1
    try:
        d = page.get_text('dict')
        for block in d.get('blocks', []):
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                bbox = fitz.Rect(line['bbox'])
                if in_table(bbox):
                    continue
                text = ''.join((sp.get('text') or '') for sp in line.get('spans', []))
                text = text.rstrip()
                if not text.strip():
                    continue
                items.append((bbox.y0, seq, 'text', text))
                seq += 1
    except Exception:  # noqa: BLE001
        pass
    items.sort(key=lambda it: (round(it[0], 1), it[1]))
    out = []
    for _y0, _seq, kind, payload in items:
        if kind == 'table':
            out.append(payload)
            out.append('')
        else:
            if _looks_like_formula(payload):
                out.append('$' + payload.strip() + '$')
            else:
                out.append(payload)
            out.append('')
    return '\n'.join(out).strip()


def pdf2md(path):
    import fitz
    doc = fitz.open(path)
    parts = []
    try:
        for page in doc:
            p = _page_to_md(page)
            if p.strip():
                parts.append(p)
    finally:
        doc.close()
    text = '\n\n'.join(parts).strip()
    if not text:
        raise ValueError('pdf 未提取到文字内容（可能是扫描件，请用 OCR）')
    return text
