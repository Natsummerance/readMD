# -*- coding: utf-8 -*-
"""万物转 md：核心格式专用解析（docx / pdf）+ MarkItDown 兜底 + 严格校验。

v2.1.1 质量升级：
- .docx：python-docx + lxml 专用解析 —— OMML 公式转 LaTeX、标题层级、表格、
  等宽字体代码块、图片引用；
- .pdf：PyMuPDF 专用解析 —— find_tables 还原边框表格 + 公式启发式；
- .xlsx/.pptx：内置 OOXML 轻量解析（工作表表格 / 幻灯片标题与段落），不依赖第三方库；
- 旧版 .xls/.ppt 与其余格式（html/csv/json/zip 等）走 MarkItDown；
- 专用解析异常时逐文件回退 MarkItDown，仍失败抛出带原因的异常。
"""

import os
import sys
import ntpath
import re
import shutil
import subprocess
import tempfile
import time
import zipfile
import logging
import html as _html
import xml.etree.ElementTree as ET
from collections import Counter

_engine = None

_MONO_FONTS = ('consolas', 'courier new', 'courier', 'menlo', 'monaco',
               'source code pro', 'cascadia code', 'fira code', 'jetbrains mono',
               'sf mono', 'liberation mono', 'dejavu sans mono')

_CODE_LANG_HINTS = (
    ('python', 'python'), ('javascript', 'javascript'), ('typescript', 'typescript'),
    ('java', 'java'), ('csharp', 'csharp'), ('c++', 'cpp'), ('cpp', 'cpp'),
    ('go', 'go'), ('rust', 'rust'), ('sql', 'sql'), ('html', 'html'),
    ('css', 'css'), ('bash', 'bash'), ('shell', 'bash'), ('json', 'json'),
    ('xml', 'xml'), ('yaml', 'yaml'), ('powershell', 'powershell'),
    ('php', 'php'), ('ruby', 'ruby'), ('vb', 'vb'), ('swift', 'swift'),
    ('kotlin', 'kotlin'), ('c#', 'csharp'), ('c', 'c'), ('cs', 'csharp'),
)

EXT_TO_LANG = {
    '.toml': 'toml', '.yaml': 'yaml', '.yml': 'yaml', '.json': 'json',
    '.json5': 'json5', '.jsonc': 'jsonc', '.ini': 'ini', '.cfg': 'ini',
    '.conf': 'ini', '.config': 'xml', '.env': 'bash', '.properties': 'properties',
    '.xml': 'xml', '.plist': 'xml', '.inf': 'ini', '.bat': 'batch',
    '.cmd': 'batch', '.ps1': 'powershell', '.psm1': 'powershell', '.sh': 'bash',
    '.bash': 'bash', '.zsh': 'bash', '.fish': 'fish', '.vbs': 'vbscript',
    '.py': 'python', '.js': 'javascript', '.mjs': 'javascript', '.cjs': 'javascript',
    '.ts': 'typescript', '.tsx': 'tsx', '.jsx': 'jsx', '.c': 'c',
    '.cpp': 'cpp', '.h': 'c', '.hpp': 'cpp', '.cc': 'cpp', '.cxx': 'cpp',
    '.cs': 'csharp', '.java': 'java', '.kt': 'kotlin', '.kts': 'kotlin',
    '.rs': 'rust', '.go': 'go', '.rb': 'ruby', '.php': 'php',
    '.swift': 'swift', '.lua': 'lua', '.r': 'r', '.m': 'objectivec',
    '.dart': 'dart', '.sql': 'sql', '.dockerfile': 'dockerfile', '.makefile': 'makefile',
    '.gradle': 'groovy', '.css': 'css',
    '.scss': 'scss', '.sass': 'sass', '.less': 'less', '.vue': 'vue',
    '.svelte': 'svelte', '.log': 'log', '.out': 'log', '.err': 'log',
    '.diff': 'diff', '.patch': 'diff', '.gitignore': 'gitignore',
    '.gitattributes': 'gitignore', '.editorconfig': 'ini', '.npmrc': 'ini',
    '.rst': 'rst', '.asciidoc': 'asciidoc', '.adoc': 'asciidoc',
    '.bib': 'bibtex', '.tex': 'latex', '.latex': 'latex',
    '.csv': 'csv', '.tsv': 'tsv',
}

# Never pass archives, executables or media through a text decoder. A
# replacement-decoded binary would look like a successful but unusable MD
# conversion; images are handled by the OCR lane instead.
_BINARY_ONLY_EXTS = frozenset({
    '.zip', '.7z', '.rar', '.tar', '.gz', '.bz2', '.xz', '.iso',
    '.exe', '.dll', '.so', '.dylib', '.bin', '.dat', '.db', '.sqlite',
    '.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tif', '.tiff',
    '.mp3', '.wav', '.mp4', '.mov', '.avi', '.mkv', '.woff', '.woff2',
})

_MATH_CHARS = set('+-*/=<>^_~%∑∫√∞≈≠±×÷∂∇∏πθαβγδφψωλμνστρΔΩΦΓΛεηζξοπς')
# 强数学算子：连字符/百分号不算（日期、电话、百分数会被密度误判成公式）
_STRONG_MATH = set('=+*/^_~<>' + '∑∫√∞≈≠±×÷∂∇∏π')


def load():
    # 惰性加载：docx / pdf 专用解析不依赖 markitdown（Win7 版未安装）
    return _engine


# ---------------------------------------------------------------- 入口

def convert(path):
    """把任意支持的文件转换为 Markdown 文本（专用解析 → MarkItDown 兜底）。"""
    text, _engine_name, error = convert_verbose(path)
    if error and not text:
        raise ValueError(error)
    return text


def csv2md(path, delimiter=None):
    """CSV / TSV 格式化转 Markdown 表格。"""
    import csv
    import io
    from . import txtmd
    text, _enc = txtmd.read_text(path)
    if not text.strip():
        return "# %s\n\n*(空文件)*" % os.path.basename(path)

    if delimiter is None:
        ext = os.path.splitext(path)[1].lower()
        if ext == '.tsv':
            delimiter = '\t'
        else:
            first_line = text.splitlines()[0] if text.splitlines() else ''
            delimiter = '\t' if '\t' in first_line and ',' not in first_line else ','

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [r for r in reader if r and any(cell.strip() for cell in r)]
    if not rows:
        return "# %s\n\n*(空表格)*" % os.path.basename(path)

    col_count = max(len(r) for r in rows)
    norm_rows = [r + [''] * (col_count - len(r)) for r in rows]

    header = norm_rows[0]
    data_rows = norm_rows[1:]

    def esc(s):
        return str(s).replace('\n', '<br>').replace('|', '\\|').strip()

    md_lines = ["# %s" % os.path.basename(path), ""]
    md_lines.append("| " + " | ".join(esc(c) for c in header) + " |")
    md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for r in data_rows:
        md_lines.append("| " + " | ".join(esc(c) for c in r) + " |")

    return "\n".join(md_lines)


def code2md(path, ext=None):
    """纯文本 / 代码 / 配置 / 日志转带高亮与信息栏的结构化 Markdown。"""
    from . import txtmd
    text, _enc = txtmd.read_text(path)
    filename = os.path.basename(path)
    if ext is None:
        ext = os.path.splitext(path)[1].lower()

    lang = EXT_TO_LANG.get(ext, '')
    lines_count = len(text.splitlines())
    size_kb = len(text.encode('utf-8')) / 1024.0

    out = [
        "# %s" % filename,
        "",
        "> **文件信息**：`%s` · %d 行 · %.1f KB" % (lang or 'plain text', lines_count, size_kb),
        "",
        "```%s" % lang,
        text,
        "```",
        ""
    ]
    return "\n".join(out)


def _rtf_to_md(path):
    """Convert the text subset of RTF without shelling out to Word/LibreOffice.

    Formatting controls that cannot be represented safely are ignored; plain
    text, paragraphs, tabs and basic bold/italic markers are preserved.
    """
    raw = open(path, 'rb').read().decode('latin-1', errors='replace')
    out, i, skip = [], 0, False
    # RTF destinations (fonttbl/colortbl/pict/...) are scoped to their
    # enclosing group.  A single boolean without a group stack permanently
    # suppresses the document body after the first font table.
    skip_stack = []
    while i < len(raw):
        ch = raw[i]
        if ch == '{':
            skip_stack.append(skip); i += 1; continue
        if ch == '}':
            skip = skip_stack.pop() if skip_stack else False
            i += 1; continue
        if ch != '\\':
            if not skip:
                out.append(ch)
            i += 1; continue
        i += 1
        if i >= len(raw):
            break
        if raw[i] in '{}\\':
            if not skip:
                out.append(raw[i])
            i += 1; continue
        if raw[i] == "'" and i + 2 < len(raw):
            try:
                if not skip:
                    out.append(bytes.fromhex(raw[i + 1:i + 3]).decode('cp1252', errors='replace'))
            except ValueError:
                pass
            i += 3; continue
        m = re.match(r'([a-zA-Z]+)(-?\d+)? ?', raw[i:])
        if not m:
            i += 1; continue
        word = m.group(1).lower()
        i += len(m.group(0))
        if word in ('fonttbl', 'colortbl', 'stylesheet', 'info', 'pict'):
            skip = True
        elif word == 'par':
            if not skip:
                out.append('\n')
        elif word == 'line':
            if not skip:
                out.append('\n')
        elif word == 'tab':
            if not skip:
                out.append('\t')
        elif word in ('b', 'i', 'ul') and not skip:
            out.append('**' if word == 'b' else '*')
    text = _html.unescape(''.join(out))
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    if not text:
        raise ValueError('rtf-empty')
    return '# %s\n\n%s\n' % (os.path.basename(path), text)


def _odt_to_md(path):
    """Extract paragraphs, headings and simple tables from an ODT package."""
    ns = {
        'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
        'table': 'urn:oasis:names:tc:opendocument:xmlns:table:1.0',
    }
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read('content.xml'))
    lines = ['# %s' % os.path.basename(path), '']
    body = root.find('.//{urn:oasis:names:tc:opendocument:xmlns:office:1.0}body')
    if body is None:
        raise ValueError('odt-empty')
    for node in body.iter():
        tag = node.tag.rsplit('}', 1)[-1] if isinstance(node.tag, str) else ''
        if tag in ('h', 'p'):
            text = ''.join(node.itertext()).strip()
            if not text:
                continue
            if tag == 'h':
                level = min(6, int(node.get('{%s}outline-level' % ns['text'], '1') or 1))
                lines.extend(['%s %s' % ('#' * level, text), ''])
            else:
                lines.extend([text, ''])
        elif tag == 'table':
            rows = []
            for row in node.findall('.//table:table-row', ns):
                cells = [' '.join(''.join(c.itertext()).split()) for c in row.findall('table:table-cell', ns)]
                if cells:
                    rows.append(cells)
            if rows:
                width = max(len(r) for r in rows)
                rows = [r + [''] * (width - len(r)) for r in rows]
                lines.append('| ' + ' | '.join(rows[0]) + ' |')
                lines.append('| ' + ' | '.join(['---'] * width) + ' |')
                lines.extend('| ' + ' | '.join(r) + ' |' for r in rows[1:])
                lines.append('')
    if len(lines) <= 2:
        raise ValueError('odt-empty')
    return '\n'.join(lines).strip() + '\n'


def _html_fragment_to_md(content):
    content = re.sub(r'<(script|style|noscript)\b[^>]*>.*?</\1>', '', content, flags=re.I | re.S)
    for level in range(6, 0, -1):
        content = re.sub(r'<h%d\b[^>]*>(.*?)</h%d>' % (level, level),
                         lambda m: '\n%s %s\n' % ('#' * level, re.sub(r'<[^>]+>', '', m.group(1))),
                         content, flags=re.I | re.S)
    content = re.sub(r'<(br|p|div|li|tr)\b[^>]*>', '\n', content, flags=re.I)
    content = re.sub(r'</(p|div|li|tr|table|ul|ol)>', '\n', content, flags=re.I)
    content = re.sub(r'<[^>]+>', '', content)
    content = _html.unescape(content)
    return re.sub(r'\n{3,}', '\n\n', content).strip()


def _epub_to_md(path):
    """Convert EPUB spine XHTML into readable Markdown text."""
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        html_names = [n for n in names if n.lower().endswith(('.xhtml', '.html', '.htm'))]
        if not html_names:
            raise ValueError('epub-no-content')
        chunks = []
        for name in sorted(html_names):
            text = archive.read(name).decode('utf-8', errors='replace')
            converted = _html_fragment_to_md(text)
            if converted:
                chunks.append(converted)
    if not chunks:
        raise ValueError('epub-empty')
    return '# %s\n\n%s\n' % (os.path.basename(path), '\n\n'.join(chunks))


_XLS_NS = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
_PPT_NS = 'http://schemas.openxmlformats.org/presentationml/2006/main'
_RID_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def _local(tag):
    return tag.rsplit('}', 1)[-1] if isinstance(tag, str) else ''


def _md_cell(text):
    return ' '.join(str(text).split()).replace('|', '\\|')


def _ooxml_column_index(ref):
    """'BC12' -> 54（0 基列号）；引用缺字母时返回 None。"""
    letters = ''.join(ch for ch in ref if ch.isalpha()).upper()
    if not letters:
        return None
    index = 0
    for ch in letters:
        index = index * 26 + (ord(ch) - 64)
    return index - 1


def _ooxml_rels(archive, names, base_dir, base_name):
    """读取包内 rels，映射 rId -> 归一化到 base_dir 下的存档路径。

    rels 文件名跟随主文件名（如 workbook.xml.rels），与 base_dir 无关。
    """
    rels = {}
    rel_path = '%s/_rels/%s.xml.rels' % (base_dir, base_name)
    if rel_path not in names:
        return rels
    root = ET.fromstring(archive.read(rel_path))
    for rel in root:
        rid = rel.get('Id')
        target = (rel.get('Target') or '').lstrip('/')
        if rid and target:
            rels[rid] = target if target.startswith('%s/' % base_dir) else '%s/%s' % (base_dir, target)
    return rels


def _xlsx_to_md(path):
    """Extract per-sheet tables from an XLSX package into Markdown tables."""
    ns = '{%s}' % _XLS_NS
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if 'xl/workbook.xml' not in names:
            raise ValueError('xlsx-empty')
        shared = []
        if 'xl/sharedStrings.xml' in names:
            sst = ET.fromstring(archive.read('xl/sharedStrings.xml'))
            for si in sst.findall('%ssi' % ns):
                shared.append(''.join(t.text or '' for t in si.iter('%st' % ns)))
        rels = _ooxml_rels(archive, names, 'xl', 'workbook')
        workbook = ET.fromstring(archive.read('xl/workbook.xml'))
        out = ['# %s' % os.path.basename(path), '']
        produced = False
        for sheet in workbook.findall('.//%ssheet' % ns):
            target = rels.get(sheet.get('{%s}id' % _RID_NS) or '')
            if not target or target not in names:
                continue
            root = ET.fromstring(archive.read(target))
            row_nodes = root.findall('.//%ssheetData/%srow' % (ns, ns))
            cells = {}
            max_col = -1
            for row_idx, row in enumerate(row_nodes):
                next_col = 0
                for cell in row.findall('%sc' % ns):
                    col = _ooxml_column_index(cell.get('r') or '')
                    if col is None:
                        col = next_col
                    next_col = col + 1
                    kind = cell.get('t') or 'n'
                    if kind == 'inlineStr':
                        inline = cell.find('%sis' % ns)
                        value = ''.join(t.text or '' for t in inline.iter('%st' % ns)) if inline is not None else ''
                    else:
                        v_el = cell.find('%sv' % ns)
                        raw = (v_el.text or '').strip() if v_el is not None else ''
                        if kind == 's':
                            try:
                                value = shared[int(raw)]
                            except (ValueError, IndexError):
                                value = ''
                        elif kind == 'b':
                            value = 'TRUE' if raw == '1' else 'FALSE' if raw == '0' else raw
                        else:
                            value = raw
                    cells[(row_idx, col)] = value
                    max_col = max(max_col, col)
            width = max_col + 1
            if width <= 0:
                continue
            rows = [[cells.get((r, c), '') for c in range(width)] for r in range(len(row_nodes))]
            while rows and not any(rows[0]):
                rows.pop(0)
            while rows and not any(rows[-1]):
                rows.pop()
            while width > 0 and not any(row[width - 1] for row in rows):
                width -= 1
            rows = [row[:width] for row in rows]
            if not rows:
                continue
            out.append('## %s' % (sheet.get('name') or 'Sheet'))
            out.append('')
            out.append('| ' + ' | '.join(_md_cell(c) for c in rows[0]) + ' |')
            out.append('| ' + ' | '.join(['---'] * width) + ' |')
            for row in rows[1:]:
                out.append('| ' + ' | '.join(_md_cell(c) for c in row) + ' |')
            out.append('')
            produced = True
        if not produced:
            raise ValueError('xlsx-empty')
        return '\n'.join(out).strip() + '\n'


def _pptx_to_md(path):
    """Extract slide titles, paragraphs and tables from a PPTX package."""
    p_ns = '{%s}' % _PPT_NS
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        if 'ppt/presentation.xml' not in names:
            raise ValueError('pptx-empty')
        rels = _ooxml_rels(archive, names, 'ppt', 'presentation')
        pres = ET.fromstring(archive.read('ppt/presentation.xml'))
        slide_rids = [s.get('{%s}id' % _RID_NS) or ''
                      for s in pres.findall('.//%ssldIdLst/%ssldId' % (p_ns, p_ns))]
        targets = [rels[rid] for rid in slide_rids if rid in rels and rels[rid] in names]
        if not targets:
            targets = sorted(n for n in names if re.match(r'ppt/slides/slide\d+\.xml$', n))
        if not targets:
            raise ValueError('pptx-empty')
        out = ['# %s' % os.path.basename(path), '']
        produced = False
        for idx, target in enumerate(targets, 1):
            root = ET.fromstring(archive.read(target))
            tree = root.find('%scSld/%sspTree' % (p_ns, p_ns))
            if tree is None:
                continue
            lines = []
            for node in tree.iter():
                tag = _local(node.tag)
                if tag == 'sp':
                    title = any(_local(el.tag) == 'ph' and (el.get('type') or '') in ('title', 'ctrTitle')
                                for el in node.iter())
                    texts = []
                    for el in node.iter():
                        if _local(el.tag) == 'txBody':
                            for para in el:
                                if _local(para.tag) == 'p':
                                    text = ' '.join(''.join(t.text or '' for t in para.iter()
                                                            if _local(t.tag) == 't').split())
                                    if text:
                                        texts.append(text)
                            break
                    if texts:
                        if title:
                            lines.extend(['## %s' % texts[0], ''])
                            texts = texts[1:]
                        lines.extend(line for text in texts for line in (text, ''))
                elif tag == 'tbl':
                    rows = []
                    for tr in node.iter():
                        if _local(tr.tag) != 'tr':
                            continue
                        cells = [_md_cell(''.join(t.text or '' for t in tc.iter() if _local(t.tag) == 't'))
                                 for tc in tr if _local(tc.tag) == 'tc']
                        if cells:
                            rows.append(cells)
                    if rows:
                        width = max(len(r) for r in rows)
                        rows = [r + [''] * (width - len(r)) for r in rows]
                        lines.append('| ' + ' | '.join(rows[0]) + ' |')
                        lines.append('| ' + ' | '.join(['---'] * width) + ' |')
                        lines.extend('| ' + ' | '.join(r) + ' |' for r in rows[1:])
                        lines.append('')
            if lines:
                if not any(line.startswith('## ') for line in lines):
                    lines.insert(0, '## Slide %d' % idx)
                out.extend(lines)
                produced = True
        if not produced:
            raise ValueError('pptx-empty')
        return '\n'.join(out).strip() + '\n'


def convert_verbose(path, form_tables=True):
    """返回 (text, engine, error)。engine: 'docx' | 'pdf' | 'csv' | 'code' | 'txtmd' | 'texmd' | 'xlsx' | 'pptx' | 'markitdown' | ''"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        try:
            return docx2md(path, form_tables=form_tables), 'docx', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    if ext == '.doc':
        text, err = doc2md(path, form_tables=form_tables)
        if err is None:
            return text, 'doc', None
        try:
            md = _markitdown_convert(path)
        except Exception as e2:  # noqa: BLE001
            return '', '', '%s（MarkItDown 兜底也失败：%s）' % (err, e2)
        # MarkItDown 对未知二进制会原样吐字节当"成功"，含 NUL 的输出视为垃圾，
        # 回落到上面的稳定错误码（真 RTF 文本不含 NUL，仍可被兜底救回）。
        if md and '\x00' not in md:
            return md, 'markitdown', None
        return '', '', err
    if ext == '.pdf':
        try:
            return pdf2md(path), 'pdf', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    if ext in ('.csv', '.tsv'):
        try:
            return csv2md(path), 'csv', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'CSV 表格转换失败：%s' % e
    if ext in ('.tex', '.latex'):
        try:
            from . import texmd
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                tex_content = f.read()
            return texmd.latex_to_md(tex_content, base_dir=os.path.dirname(os.path.abspath(path))), 'texmd', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'LaTeX 转换失败：%s' % e
    if ext in ('.txt', '.text'):
        try:
            from . import txtmd
            text, _enc = txtmd.read_text(path)
            md, tstats = txtmd.to_markdown(text)
            return md, 'txtmd', None
        except Exception as e:  # noqa: BLE001
            return '', '', '文本转换失败：%s' % e
    if ext == '.rtf':
        try:
            return _rtf_to_md(path), 'rtf', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'RTF 转换失败：%s' % e
    if ext == '.odt':
        try:
            return _odt_to_md(path), 'odt', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'ODT 转换失败：%s' % e
    if ext == '.epub':
        try:
            return _epub_to_md(path), 'epub', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'EPUB 转换失败：%s' % e
    if ext == '.xlsx':
        try:
            return _xlsx_to_md(path), 'xlsx', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    if ext == '.pptx':
        try:
            return _pptx_to_md(path), 'pptx', None
        except Exception as e:  # noqa: BLE001
            try:
                return _markitdown_convert(path), 'markitdown', None
            except Exception as e2:  # noqa: BLE001
                return '', '', '%s（MarkItDown 兜底也失败：%s）' % (e, e2)
    if ext in ('.xls', '.ppt'):
        # 旧版二进制 Office 格式没有内置解析器；仅在装有 MarkItDown 时可转换。
        try:
            return _markitdown_convert(path), 'markitdown', None
        except Exception as e:  # noqa: BLE001
            return '', '', 'legacy-office：旧版 %s 二进制格式需安装 MarkItDown 才能转换（%s）' % (ext, e)
    if ext in EXT_TO_LANG:
        try:
            return code2md(path, ext), 'code', None
        except Exception as e:  # noqa: BLE001
            return '', '', '代码/配置格式化转换失败：%s' % e

    # Binary payloads must fail explicitly instead of becoming unreadable
    # replacement-decoded Markdown. Images are handled by the OCR lane.
    if ext in _BINARY_ONLY_EXTS or _looks_binary(path):
        return '', '', 'unsupported_format'

    # 兜底：先尝试 MarkItDown，若失败尝试通用文本读取 code2md 兜底
    try:
        return _markitdown_convert(path), 'markitdown', None
    except Exception as e:  # noqa: BLE001
        try:
            return code2md(path, ext), 'code_fallback', None
        except Exception:
            return '', '', str(e)



OLE2_MAGIC = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'


def _find_soffice():
    """定位 LibreOffice soffice 可执行文件；找不到返回 None。"""
    exe = shutil.which('soffice') or shutil.which('soffice.exe')
    if exe:
        return exe
    candidates = (
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        '/usr/bin/soffice',
        '/usr/local/bin/soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    )
    for c in candidates:
        if os.path.isfile(c):
            return c
    return None


_WORD_RETRY_HRESULTS = frozenset({
    -2147418111,  # RPC_E_CALL_REJECTED：调用被呼叫方拒绝（Word 冷启动/忙）
    -2147417846,  # RPC_E_SERVERCALL_RETRYLATER：服务器忙，稍后重试
    -2146959355,  # CO_E_SERVER_EXEC_FAILURE：服务器启动失败（冷启动竞争）
})


def _word_com_retry(fn, attempts=4, delays=(0.5, 1.0, 2.0)):
    """对 Word COM 暂态故障（调用被拒 / 服务器忙 / 启动失败）做有界重试。

    其余异常立即抛出；重试耗尽后抛出最后一次异常。
    """
    for attempt in range(attempts):
        try:
            return fn()
        except Exception as e:  # noqa: BLE001
            hresult = getattr(e, 'hresult', None)
            if hresult is None and e.args and isinstance(e.args[0], int):
                hresult = e.args[0]
            if hresult not in _WORD_RETRY_HRESULTS or attempt == attempts - 1:
                raise
            time.sleep(delays[min(attempt, len(delays) - 1)])


def _word_com_process_worker(src, out_dir, queue):
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        queue.put((False, None))
        return

    pythoncom.CoInitialize()
    word = None
    try:
        word = _word_com_retry(
            lambda: win32com.client.DispatchEx('Word.Application'))
        word.Visible = False
        word.DisplayAlerts = 0

        abs_src = os.path.abspath(src)
        abs_out_dir = os.path.abspath(out_dir)
        out = os.path.join(abs_out_dir, os.path.splitext(os.path.basename(src))[0] + '.docx')

        def _open_and_save():
            try:
                doc = word.Documents.Open(
                    abs_src,
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    Visible=False
                )
            except TypeError:
                doc = word.Documents.Open(abs_src, ReadOnly=True)
            try:
                doc.SaveAs2(out, FileFormat=16)  # wdFormatXMLDocument
            finally:
                doc.Close(False)
            return out

        res = _word_com_retry(_open_and_save)
        queue.put((True, res if os.path.isfile(res) else None))
    except Exception as e:
        queue.put((False, str(e)))
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _doc2docx_word_com(src, out_dir, timeout=30):
    """用本机 Word COM 把 .doc 转成 .docx；带进程级超时强杀保护与弹窗压制。"""
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return None

    # When win32com is mocked in unit tests, run in-process to allow assertions on mock objects
    is_mock = getattr(sys.modules.get('win32com'), '__file__', None) is None
    if is_mock:
        class _DirectQueue:
            def __init__(self):
                self.val = None
            def put(self, v):
                self.val = v
        dq = _DirectQueue()
        _word_com_process_worker(src, out_dir, dq)
        if dq.val and dq.val[0]:
            return dq.val[1]
        return None

    import multiprocessing
    q = multiprocessing.Queue()
    p = multiprocessing.Process(target=_word_com_process_worker, args=(src, out_dir, q))
    p.start()
    p.join(timeout=timeout)
    if p.is_alive():
        logging.warning("Word COM conversion timed out after %ds for %s; terminating worker process", timeout, src)
        # On Windows Word is a separate COM child.  Terminating only the
        # Python worker can leave WINWORD.EXE behind and keep the file locked.
        if os.name == 'nt' and getattr(p, 'pid', None):
            try:
                subprocess.run(
                    ['taskkill', '/PID', str(p.pid), '/T', '/F'],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                    check=False, timeout=5,
                )
            except Exception:
                pass
        try:
            p.terminate()
        finally:
            p.join(timeout=3)
        return None

    # ``Queue.empty()`` is racy because the feeder thread may not have flushed
    # its last item even after the worker exited.  A bounded get is reliable
    # and keeps a broken worker from blocking the caller forever.
    try:
        ok, res = q.get(timeout=2)
        if ok and res and os.path.isfile(res):
            return res
    except Exception:
        pass
    finally:
        try:
            q.close()
            q.join_thread()
        except Exception:
            pass
    return None


def _doc2docx_soffice(src, out_dir):
    """用 LibreOffice headless 把 .doc 转成 .docx；未安装时返回 None。"""
    exe = _find_soffice()
    if not exe:
        return None
    subprocess.run(
        [exe, '--headless', '--convert-to', 'docx', '--outdir', out_dir,
         os.path.abspath(src)],
        timeout=180, capture_output=True, check=True)
    out = os.path.join(out_dir, os.path.splitext(os.path.basename(src))[0] + '.docx')
    return out if os.path.isfile(out) else None


def doc2md(path, form_tables=True):
    """Word 97-2003 .doc 专用解析：magic bytes 校验 → Word COM → soffice 兜底。

    返回 (text, error)；error 为 None 表示成功。
    """
    try:
        with open(path, 'rb') as f:
            magic = f.read(8)
    except OSError as e:
        return '', 'doc-read-failed：%s' % e
    if magic != OLE2_MAGIC:
        return '', 'doc-not-ole2：不是有效的 Word 97-2003 文档（magic bytes 不匹配）'
    last = None
    with tempfile.TemporaryDirectory(prefix='readmd-doc-') as td:
        for converter in (_doc2docx_word_com, _doc2docx_soffice):
            try:
                out = converter(path, td)
            except Exception as e:  # noqa: BLE001
                last = e
                continue
            if out and os.path.isfile(out):
                return docx2md(out, form_tables=form_tables), None
    if last is not None:
        return '', 'doc-convert-failed：转换引擎执行失败（%s）' % last
    return '', 'doc-no-engine：本机未找到可用的 DOC 转换引擎（需要安装 Microsoft Word 或 LibreOffice）'


def _markitdown_convert(path):
    global _engine
    if _engine is None:
        try:
            from markitdown import MarkItDown
        except Exception as e:  # noqa: BLE001
            raise ImportError('MarkItDown 未安装（%s），本格式无法转换' % e)
        _engine = MarkItDown()
    result = _engine.convert(path)
    return (result.text_content or '').strip()


def _looks_binary(path):
    """Cheap guard against replacement-decoding arbitrary binary payloads."""
    try:
        with open(path, 'rb') as handle:
            sample = handle.read(8192)
    except OSError:
        return False
    if not sample or b'\x00' in sample:
        return bool(sample and b'\x00' in sample)
    try:
        decoded = sample.decode('utf-8')
    except UnicodeDecodeError:
        # Legacy encodings are still eligible for the text readers.
        return False
    return decoded.count('\ufffd') / max(1, len(decoded)) > 0.02


# ---------------------------------------------------------------- docx 专用

_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

_UNICODE_MATH_TO_LATEX = {
    'α': r'\alpha', 'β': r'\beta', 'γ': r'\gamma', 'δ': r'\delta',
    'ϵ': r'\epsilon', 'ε': r'\varepsilon', 'ζ': r'\zeta', 'η': r'\eta',
    'θ': r'\theta', 'ϑ': r'\vartheta', 'ι': r'\iota', 'κ': r'\kappa',
    'λ': r'\lambda', 'μ': r'\mu', 'ν': r'\nu', 'ξ': r'\xi',
    'π': r'\pi', 'ϖ': r'\varpi', 'ρ': r'\rho', 'ϱ': r'\varrho',
    'σ': r'\sigma', 'ς': r'\varsigma', 'τ': r'\tau', 'υ': r'\upsilon',
    'ϕ': r'\phi', 'φ': r'\varphi', 'χ': r'\chi', 'ψ': r'\psi', 'ω': r'\omega',
    'Γ': r'\Gamma', 'Δ': r'\Delta', 'Θ': r'\Theta', 'Λ': r'\Lambda',
    'Ξ': r'\Xi', 'Π': r'\Pi', 'Σ': r'\Sigma', 'Υ': r'\Upsilon',
    'Φ': r'\Phi', 'Ψ': r'\Psi', 'Ω': r'\Omega',
    '±': r'\pm', '∓': r'\mp', '×': r'\times', '÷': r'\div', '·': r'\cdot',
    '∗': r'\ast', '⋆': r'\star', '∘': r'\circ', '∙': r'\bullet',
    '≤': r'\leq', '≥': r'\geq', '≠': r'\neq', '≈': r'\approx', '≡': r'\equiv',
    '∼': r'\sim', '≃': r'\simeq', '≅': r'\cong', '∝': r'\propto',
    '≪': r'\ll', '≫': r'\gg',
    '→': r'\to', '←': r'\leftarrow', '⇒': r'\Rightarrow', '⇐': r'\Leftarrow',
    '↔': r'\leftrightarrow', '⇔': r'\Leftrightarrow', '↦': r'\mapsto',
    '↑': r'\uparrow', '↓': r'\downarrow',
    '∞': r'\infty', '∂': r'\partial', '∇': r'\nabla', '′': r"'", 'ℏ': r'\hbar',
    '∈': r'\in', '∉': r'\notin', '⊂': r'\subset', '⊆': r'\subseteq',
    '⊃': r'\supset', '⊇': r'\supseteq', '∩': r'\cap', '∪': r'\cup',
    '∖': r'\setminus', '∀': r'\forall', '∃': r'\exists', '¬': r'\neg',
    '∧': r'\land', '∨': r'\lor', '∅': r'\emptyset',
    '…': r'\ldots', '⋯': r'\cdots', '⋮': r'\vdots', '⋱': r'\ddots',
    '∠': r'\angle', '⊥': r'\perp', '∥': r'\parallel',
    '⟨': r'\langle', '⟩': r'\rangle',
    '⊗': r'\otimes', '⊕': r'\oplus', '⊙': r'\odot',
    '∑': r'\sum', '∏': r'\prod', '∐': r'\coprod',
    '∫': r'\int', '∬': r'\iint', '∭': r'\iiint', '∮': r'\oint',
    '⋂': r'\bigcap', '⋃': r'\bigcup'
}


def _mq(tag):
    return '{%s}%s' % (_M_NS, tag)


def _omml_to_latex(el):
    """递归把 OMML 元素转为高质量 LaTeX 表达式。"""
    if el is None:
        return ''
    tag = el.tag
    if not isinstance(tag, str) or not tag.startswith('{'):
        return (el.text or '')
    local = tag.split('}')[1]

    def kids(e):
        if e is None:
            return ''
        return ''.join(_omml_to_latex(c) for c in e)

    if local == 't':
        txt = (el.text or '').replace('\u200b', '').replace('\u2061', '')
        # 替换 Unicode 数学符号
        out_chars = []
        for ch in txt:
            if ch in _UNICODE_MATH_TO_LATEX:
                out_chars.append(_UNICODE_MATH_TO_LATEX[ch] + ' ')
            else:
                out_chars.append(ch)
        return ''.join(out_chars)

    if local == 'r':
        rpr = el.find(_mq('rPr'))
        inner = kids(el).strip()
        if rpr is not None and inner:
            if rpr.find(_mq('nor')) is not None:
                return r'\text{%s}' % inner
            if rpr.find(_mq('b')) is not None:
                return r'\mathbf{%s}' % inner
            i_el = rpr.find(_mq('i'))
            if i_el is not None and (i_el.get(_mq('val')) == 'off' or i_el.get('val') == 'off'):
                return r'\mathrm{%s}' % inner
        return kids(el)

    if local in ('oMath', 'oMathPara', 'e', 'num', 'den', 'sub',
                 'sup', 'deg', 'chr', 'fName', 'delim',
                 'phant'):
        return kids(el)

    if local == 'f':  # 分数
        num = el.find(_mq('num'))
        den = el.find(_mq('den'))
        fPr = el.find(_mq('fPr'))
        if fPr is not None:
            t_type = fPr.find(_mq('type'))
            if t_type is not None and (t_type.get(_mq('val')) == 'noBar' or t_type.get('val') == 'noBar'):
                return r'\binom{%s}{%s}' % (kids(num).strip(), kids(den).strip())
        return r'\frac{%s}{%s}' % (kids(num).strip(), kids(den).strip())

    if local == 'sSup':
        base = el.find(_mq('e'))
        sup = el.find(_mq('sup'))
        return '{%s}^{%s}' % (kids(base).strip(), kids(sup).strip())

    if local == 'sSub':
        base = el.find(_mq('e'))
        sub = el.find(_mq('sub'))
        return '{%s}_{%s}' % (kids(base).strip(), kids(sub).strip())

    if local == 'sSubSup':
        base = el.find(_mq('e'))
        sub = el.find(_mq('sub'))
        sup = el.find(_mq('sup'))
        return '{%s}_{%s}^{%s}' % (kids(base).strip(), kids(sub).strip(), kids(sup).strip())

    if local == 'rad':  # 根式
        deg = el.find(_mq('deg'))
        e_el = el.find(_mq('e'))
        inner = kids(e_el).strip()
        d = kids(deg).strip()
        if d:
            return r'\sqrt[%s]{%s}' % (d, inner)
        return r'\sqrt{%s}' % inner

    if local == 'nary':  # 求和/积分/连乘
        chr_el = el.find(_mq('naryPr'))
        c = ''
        if chr_el is not None:
            c_node = chr_el.find(_mq('chr'))
            if c_node is not None:
                c = c_node.get(_mq('val'), c_node.get('val', ''))
        if not c:
            c_node = el.find(_mq('chr'))
            if c_node is not None:
                c = c_node.get(_mq('val'), c_node.get('val', '')) or kids(c_node)
        c = c.strip()
        op = {'\u2211': r'\sum', '\u222b': r'\int', '\u220f': r'\prod',
              '\u222e': r'\oint', '\u22c2': r'\bigcap', '\u22c3': r'\bigcup',
              '∑': r'\sum', '∫': r'\int', '∏': r'\prod', '∮': r'\oint'}.get(c, c or r'\int' if 'int' in c else r'\sum')
        sub = el.find(_mq('sub'))
        sup = el.find(_mq('sup'))
        e_el = el.find(_mq('e'))
        s, p = kids(sub).strip(), kids(sup).strip()
        inner = kids(e_el).strip()
        subsup = ''
        if s and p:
            subsup = '_{%s}^{%s}' % (s, p)
        elif s:
            subsup = '_{%s}' % s
        elif p:
            subsup = '^{%s}' % p

        if inner:
            return '%s%s %s' % (op, subsup, inner)
        return '%s%s' % (op, subsup)

    if local == 'd':  # 定界符 / 括号组
        dpr = el.find(_mq('dPr'))
        beg = '('
        end = ')'
        if dpr is not None:
            beg_el = dpr.find(_mq('begChr'))
            end_el = dpr.find(_mq('endChr'))
            if beg_el is not None:
                beg = beg_el.get(_mq('val'), beg_el.get('val', '('))
            if end_el is not None:
                end = end_el.get(_mq('val'), end_el.get('val', ')'))

        delim_map = {
            '(': (r'\left(', r'\right)'),
            '[': (r'\left[', r'\right]'),
            '{': (r'\left\{', r'\right\}'),
            '|': (r'\left|', r'\right|'),
            '‖': (r'\left\|', r'\right\|'),
            '⟨': (r'\left\langle', r'\right\rangle'),
            '<': (r'\left\langle', r'\right\rangle'),
            '': (r'\left.', r'\right.')
        }
        l_beg, _ = delim_map.get(beg, (r'\left%s' % beg if beg else r'\left.', ''))
        _, r_end = delim_map.get(end, ('', r'\right%s' % end if end else r'\right.'))

        e_el = el.find(_mq('e'))
        inner = kids(e_el).strip()

        # 如果内部为矩阵，转换为标准 pmatrix / bmatrix / vmatrix
        if inner.startswith(r'\begin{matrix}') and inner.endswith(r'\end{matrix}'):
            mat_body = inner[len(r'\begin{matrix}'):-len(r'\end{matrix}')].strip()
            if beg == '(' and end == ')':
                return r'\begin{pmatrix} %s \end{pmatrix}' % mat_body
            if beg == '[' and end == ']':
                return r'\begin{bmatrix} %s \end{bmatrix}' % mat_body
            if beg == '{' and end == '}':
                return r'\begin{Bmatrix} %s \end{Bmatrix}' % mat_body
            if beg == '|' and end == '|':
                return r'\begin{vmatrix} %s \end{vmatrix}' % mat_body

        return '%s %s %s' % (l_beg, inner, r_end)

    if local == 'm':  # 矩阵
        rows = el.findall(_mq('mr'))
        if rows:
            row_strs = []
            for r in rows:
                cells = r.findall(_mq('e'))
                row_strs.append(' & '.join(kids(c).strip() for c in cells))
            return r'\begin{matrix} %s \end{matrix}' % r' \\ '.join(row_strs)
        return kids(el)

    if local == 'eqArr':  # 方程组 / 多行对齐
        rows = el.findall(_mq('e'))
        if rows:
            row_strs = [kids(r).strip() for r in rows if kids(r).strip()]
            return r'\begin{aligned} %s \end{aligned}' % r' \\ '.join(row_strs)
        return kids(el)

    if local == 'limLow':  # 下标极限 / 最小值
        e_el = el.find(_mq('e'))
        lim_el = el.find(_mq('lim'))
        e_txt = kids(e_el).strip()
        l_txt = kids(lim_el).strip()
        if e_txt.lower() in ('lim', 'max', 'min', 'inf', 'sup', 'det', 'gcd'):
            return r'\%s_{%s}' % (e_txt.lower(), l_txt)
        return '{%s}_{%s}' % (e_txt, l_txt)

    if local == 'limUpp':  # 上标极限
        e_el = el.find(_mq('e'))
        lim_el = el.find(_mq('lim'))
        e_txt = kids(e_el).strip()
        l_txt = kids(lim_el).strip()
        if e_txt.lower() in ('lim', 'max', 'min', 'inf', 'sup', 'det', 'gcd'):
            return r'\%s^{%s}' % (e_txt.lower(), l_txt)
        return '{%s}^{%s}' % (e_txt, l_txt)

    if local in ('box', 'borderBox'):  # 框选
        e_el = el.find(_mq('e'))
        return r'\boxed{%s}' % kids(e_el).strip()

    if local == 'func':  # 函数
        fname = el.find(_mq('fName'))
        e_el = el.find(_mq('e'))
        fn_str = kids(fname).strip()
        in_str = kids(e_el).strip()
        known_funcs = {'sin', 'cos', 'tan', 'cot', 'sec', 'csc', 'ln', 'log', 'lg', 'exp', 'arcsin', 'arccos', 'arctan', 'sinh', 'cosh', 'tanh'}
        if fn_str.lower() in known_funcs:
            fn_str = r'\%s' % fn_str.lower()
        if in_str:
            return r'%s(%s)' % (fn_str, in_str)
        return fn_str

    if local == 'acc':  # 帽子/箭头/重音
        acc_pr = el.find(_mq('accPr'))
        a = '^'  # 默认 hat
        if acc_pr is not None:
            chr_el = acc_pr.find(_mq('chr'))
            if chr_el is not None:
                a = chr_el.get(_mq('val'), chr_el.get('val', '^'))
        e_el = el.find(_mq('e'))
        inner = kids(e_el).strip()
        marks = {'\u02c6': r'\hat', '^': r'\hat', '\u00af': r'\bar', '¯': r'\bar',
                 '\u2192': r'\vec', '→': r'\vec', '\u02dc': r'\tilde', '~': r'\tilde',
                 '\u0307': r'\dot', '˙': r'\dot', '\u0308': r'\ddot', '¨': r'\ddot',
                 'ˇ': r'\check', '´': r'\acute', '`': r'\grave'}
        cmd = marks.get(a, r'\hat')
        return (cmd + '{%s}' % inner) if cmd else inner

    if local == 'bar':  # 上下横线
        e_el = el.find(_mq('e'))
        return r'\overline{%s}' % kids(e_el).strip()

    if local == 'groupChr':  # 括号/花括号
        chr_el = el.find(_mq('chr'))
        e_el = el.find(_mq('e'))
        c = kids(chr_el).strip()
        pairs = {'{': r'\left\{ %s \right\}', '}': r'\left\{ %s \right\}',
                 '[': r'\left[ %s \right]', ']': r'\left[ %s \right]',
                 '(': r'\left( %s \right)', ')': r'\left( %s \right)',
                 '|': r'\left| %s \right|'}
        if c in pairs:
            return pairs[c] % kids(e_el).strip()
        return kids(e_el).strip()

    if local in ('rPr', 'ctrlPr', 'argPr', 'eqArrPr', 'naryPr', 'sSupPr',
                 'sSubPr', 'sSubSupPr', 'radPr', 'fPr', 'accPr', 'barPr',
                 'delimPr', 'funcPr', 'limLowPr', 'limUppPr', 'groupChrPr',
                 'phantPr', 'boxPr', 'borderBoxPr', 'mathPr', 'wrapPr',
                 'intLim', 'naryLim', 'subHide', 'supHide', 'mPr', 'mrPr'):
        return ''

    # 未知标签：取子节点文本兜底
    return kids(el)


def _run_font_lower(r):
    try:
        name = (r.font.name or '') or ''
    except Exception:  # noqa: BLE001
        name = ''
    if not name:
        rPr = r._r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            rf = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rf is not None:
                name = (rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                        or rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi')
                        or '')
    return (name or '').lower()


def _para_has_mono(p):
    return any(_run_font_lower(r) in _MONO_FONTS for r in p.runs if (r.text or '').strip())


def _para_plain(p):
    return ''.join((r.text or '') for r in p.runs)





def _lang_hint(text):
    low = (text or '').lower()
    for key, lang in _CODE_LANG_HINTS:
        if key in low:
            return lang
    return ''


def _para_inline_with_math(p, doc=None):
    """按段落子节点物理先后顺序提取文字、格式、超链接与 OMML 公式。"""
    from docx.text.run import Run
    parts = []

    for child in p._p:
        tag = child.tag
        if not isinstance(tag, str):
            continue
        local = tag.split('}')[-1]

        if local == 'r':
            r = Run(child, p)
            t = r.text or ''
            if not t:
                continue
            if _run_font_lower(r) in _MONO_FONTS:
                t = '`' + t + '`'
            if r.bold:
                t = '**' + t + '**'
            if r.italic:
                t = '*' + t + '*'
            parts.append(t)
        elif local == 'hyperlink':
            r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            url = ''
            if r_id and doc is not None and hasattr(doc, 'part') and r_id in doc.part.rels:
                try:
                    url = doc.part.rels[r_id].target_ref
                except Exception:
                    url = ''
            link_text = ''.join((c.text or '') for c in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if url and link_text:
                parts.append('[%s](%s)' % (link_text, url))
            elif link_text:
                parts.append(link_text)
        elif local == 'oMath':
            latex = _omml_to_latex(child).strip()
            if latex:
                parts.append('$%s$' % latex)
        elif local == 'oMathPara':
            for om in child.findall(_mq('oMath')):
                latex = _omml_to_latex(om).strip()
                if latex:
                    parts.append('$$%s$$' % latex)

    res = ''.join(parts).strip()
    # 如果段落内只有一个 $...$，且包含复杂结构，提升为独立公式块 $$...$$
    if res.startswith('$') and res.endswith('$') and not res.startswith('$$') and not res.endswith('$$'):
        inner_m = res[1:-1].strip()
        if len(inner_m) >= 60 or r'\begin{' in inner_m or r'\frac' in inner_m or r'\sum' in inner_m or r'\int' in inner_m or r'\aligned' in inner_m:
            res = '$$%s$$' % inner_m
    return res


def _form_table_items(rows):
    """两列短文本表 → key-value 列表项；不满足启发式条件返回空列表。"""
    if len(rows) < 2:
        return []
    for r in rows:
        if len(r) != 2 or not r[0] or len(r[0]) > 30 or len(r[1]) > 30:
            return []
        if r[0].isdigit():
            return []
        for s in r:
            if '$' in s or '](' in s or s.startswith('`'):
                return []
    return ['- **%s**: %s' % (r[0], r[1]) if r[1] else '- **%s**' % r[0] for r in rows]


def _table_to_md(tbl, doc=None, form_tables=True):
    """w:tbl → 规整管道表（展开 gridSpan/vMerge 合并单元格，保持列对齐）。"""
    from docx.table import _Cell
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    rows = []
    restart_texts = {}
    for tr in tbl._tbl.findall(ns + 'tr'):
        parsed = []
        col = 0
        for tc in tr.findall(ns + 'tc'):
            c = _Cell(tc, tbl)
            txt = '\n'.join(_para_inline_with_math(pp, doc) for pp in c.paragraphs)
            txt = txt.replace('\n', ' ').replace('|', '\\|').strip()
            span = 1
            vmerge = ''
            tc_pr = tc.find(ns + 'tcPr')
            if tc_pr is not None:
                gs = tc_pr.find(ns + 'gridSpan')
                if gs is not None:
                    try:
                        span = max(1, int(gs.get(ns + 'val') or gs.get('val') or '1'))
                    except (TypeError, ValueError):
                        span = 1
                vm = tc_pr.find(ns + 'vMerge')
                if vm is not None:
                    vmerge = (vm.get(ns + 'val') or vm.get('val') or 'continue').lower()
                    if vmerge not in ('continue', 'restart'):
                        vmerge = 'continue'
            parsed.append((txt, span, vmerge, col))
            col += span
        cells = []
        for txt, span, vmerge, col in parsed:
            if vmerge == 'continue':
                txt = restart_texts.get(col, txt)
            else:
                restart_texts[col] = txt
            cells.append(txt)
            cells.extend([''] * (span - 1))
        rows.append(cells)
    if not rows:
        return ''
    ncol = max(len(r) for r in rows)
    rows = [r + [''] * (ncol - len(r)) for r in rows]
    if form_tables and ncol == 2:
        items = _form_table_items(rows)
        if items:
            return '\n'.join(items)
    out = ['| ' + ' | '.join(rows[0]) + ' |',
           '| ' + ' | '.join(['---'] * ncol) + ' |']
    for r in rows[1:]:
        out.append('| ' + ' | '.join(r) + ' |')
    return '\n'.join(out)


def _para_list_info(p):
    """返回 (is_list, ilvl, ordered)：段落级 numPr 或样式链推断列表层级与有序性。"""
    from docx.oxml.ns import qn

    def _val(el):
        if el is None:
            return None
        return el.get(qn('w:val'), el.get('val'))

    def _ilvl(numpr):
        try:
            return max(0, min(5, int(_val(numpr.find(qn('w:ilvl'))) or 0)))
        except (TypeError, ValueError):
            return 0

    ppr = p._p.find(qn('w:pPr'))
    if ppr is not None:
        numpr = ppr.find(qn('w:numPr'))
        if numpr is not None:
            return True, _ilvl(numpr), False
    try:
        st = p.style
        depth = 0
        while st is not None and depth < 8:
            spPr = st.element.find(qn('w:pPr'))
            s_numpr = spPr.find(qn('w:numPr')) if spPr is not None else None
            if 'list number' in (st.name or '').lower():
                return True, (_ilvl(s_numpr) if s_numpr is not None else 0), True
            if s_numpr is not None:
                return True, _ilvl(s_numpr), False
            st = st.base_style
            depth += 1
    except Exception:  # noqa: BLE001
        pass
    return False, 0, False


def docx2md(path, form_tables=True):
    """docx → Markdown：OMML 公式→LaTeX、标题层级、表格、等宽字体代码块。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    base_dir = os.path.dirname(os.path.abspath(path))
    lines = []
    code_buf = []
    code_lang = ''
    ordered_n = 0

    def flush_code():
        nonlocal code_buf, code_lang
        if code_buf:
            lines.append('```' + code_lang)
            lines.extend(code_buf)
            lines.append('```')
            lines.append('')
            code_buf, code_lang = [], ''

    def handle_para(p):
        nonlocal code_buf, code_lang, ordered_n
        style_name = ''
        try:
            style_name = (p.style.name or '') if p.style is not None else ''
        except Exception:  # noqa: BLE001
            style_name = ''
        if style_name and style_name.lower().startswith('heading'):
            flush_code()
            ordered_n = 0
            try:
                level = int(''.join(ch for ch in style_name if ch.isdigit()) or '1')
            except ValueError:
                level = 1
            level = max(1, min(6, level))
            txt = _para_inline_with_math(p, doc).strip()
            if txt:
                lines.append('#' * level + ' ' + txt)
                lines.append('')
            return
        if style_name and style_name.lower() == 'title':
            flush_code()
            ordered_n = 0
            txt = _para_inline_with_math(p, doc).strip()
            if txt:
                lines.append('# ' + txt)
                lines.append('')
            return
        if _para_has_mono(p) and not p._p.findall('.//' + _mq('oMath')):
            txt = _para_plain(p).strip()
            if txt:
                if not code_lang:
                    code_lang = _lang_hint(txt)
                code_buf.append(txt)
            return
        flush_code()
        txt = _para_inline_with_math(p, doc).strip()
        if not txt:
            return
        # 列表：段落级或样式级 numPr → 缩进列表项（ilvl 控制层级，有序列表跨段计数）
        is_list, ilvl, ordered = _para_list_info(p)
        if is_list:
            if ordered:
                ordered_n += 1
                lines.append('   ' * ilvl + '%d. %s' % (ordered_n, txt))
            else:
                ordered_n = 0
                lines.append('  ' * ilvl + '- ' + txt)
            lines.append('')
            return
        ordered_n = 0
        lines.append(txt)
        lines.append('')

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            handle_para(Paragraph(child, doc))
        elif child.tag == qn('w:tbl'):
            flush_code()
            md = _table_to_md(Table(child, doc), doc, form_tables)
            if md:
                lines.append(md)
                lines.append('')
        elif child.tag == qn('w:sectPr'):
            continue
    flush_code()
    text = '\n'.join(lines).strip()
    if not text:
        raise ValueError('docx 未提取到文字内容')
    return text


# ---------------------------------------------------------------- pdf 专用

_MONO_FONTS = {'courier', 'consolas', 'monaco', 'menlo', 'firacode', 'sourcecodepro', 'mono', 'monospace', 'cascadia', 'fira'}
_BULLET_PREFIX_RE = re.compile(r'^[ \t]*[\u2022\u25cf\u25aa\u25cb\u25c6\u25c7\u25b6\u25c0\u2023\u2043\u25e6\u00b7\ufffd\x95\-*+]\s*')
_ORDERED_PREFIX_RE = re.compile(r'^[ \t]*(?:\d+[\.\)]|\(\d+\)|\[\d+\])\s*')


def _is_mono_font(font_name: str) -> bool:
    if not font_name:
        return False
    fn = font_name.lower().replace('-', '').replace('_', '').replace(' ', '')
    return any(m in fn for m in _MONO_FONTS)


def _is_cjk(ch: str) -> bool:
    if not ch:
        return False
    cp = ord(ch)
    return 0x4e00 <= cp <= 0x9fff or 0x3400 <= cp <= 0x4dbf or 0x3000 <= cp <= 0x303f or 0xff00 <= cp <= 0xffef


def _join_lines(l1: str, l2: str) -> str:
    """智能断行重组（区分中英文与连字符）。"""
    if not l1:
        return l2
    if not l2:
        return l1
    if l1.endswith('-') and len(l1) > 1 and l1[-2].isalnum() and l2[0].isalnum():
        return l1[:-1] + l2
    if _is_cjk(l1[-1]) or _is_cjk(l2[0]):
        return l1 + l2
    return l1 + ' ' + l2


def _data_to_md(data):
    """表格二维数组 → 管道表。"""
    if not data:
        return ''
    ncol = max(len(r) for r in data)
    data = [r + [''] * (ncol - len(r)) for r in data]

    def clean(cell):
        s = (cell or '') if isinstance(cell, str) else str(cell or '')
        s = s.replace('\n', ' ').replace('\r', ' ').replace('|', '\\|').strip()
        return s

    out = ['| ' + ' | '.join(clean(c) for c in data[0]) + ' |',
           '| ' + ' | '.join(['---'] * ncol) + ' |']
    for r in data[1:]:
        out.append('| ' + ' | '.join(clean(c) for c in r) + ' |')
    return '\n'.join(out)


def _pipe_cells(line):
    s = line.strip()
    return [c.strip() for c in s[1:-1].split('|')]


def _is_table_line(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|')


def _is_sep_row(line):
    cs = _pipe_cells(line)
    return bool(cs) and all(c == '---' for c in cs)


def _merge_split_tables(md):
    """合并被分页打断的相邻同列数管道表（仅限 pdf2md 输出）。"""
    lines = md.split('\n')
    out = []
    i = 0
    n = len(lines)
    in_fence = False
    while i < n:
        line = lines[i]
        if line.strip().startswith('```'):
            in_fence = not in_fence
            out.append(line)
            i += 1
            continue
        if in_fence or not _is_table_line(line):
            out.append(line)
            i += 1
            continue
        block = []
        while i < n and _is_table_line(lines[i]):
            block.append(lines[i])
            i += 1
        while True:
            j = i
            while j < n and not lines[j].strip():
                j += 1
            k = j
            while k < n and _is_table_line(lines[k]):
                k += 1
            if j >= k:
                break
            frag = lines[j:k]
            if len(_pipe_cells(frag[0])) != len(_pipe_cells(block[0])):
                break
            if _pipe_cells(frag[0]) == _pipe_cells(block[0]):
                # 重复表头：去片段首行与其中的分隔行
                for l in frag[1:]:
                    if not _is_sep_row(l):
                        block.append(l)
            else:
                # 片段首行是数据被误升格为表头：整段保留，仅去伪分隔行
                for idx, l in enumerate(frag):
                    if idx > 0 and _is_sep_row(l):
                        continue
                    block.append(l)
            i = k
        out.extend(block)
    return '\n'.join(out)


def _looks_like_formula(line):
    s = line.strip()
    if not s or len(s) < 2 or len(s) > 160:
        return False
    if any('\u4e00' <= ch <= '\u9fff' for ch in s):
        return False
    if not any(ch in _STRONG_MATH for ch in s):
        return False
    sig = sum(1 for ch in s if ch in _MATH_CHARS or ch.isdigit() or ch in '()[]{},.')
    return sig / max(len(s), 1) >= 0.3


def _accept_text_table(rows):
    """text 策略候选的行列门槛：≥2 列，且 ≥2 行各有 ≥2 个非空单元格。"""
    if not rows:
        return False
    if max(len(r) for r in rows) < 2:
        return False
    multi = 0
    for r in rows:
        if sum(1 for c in r if (c or '').strip()) >= 2:
            multi += 1
    return multi >= 2


def _retry_text_strategy_tables(page):
    """无边框表格兜底：默认策略找不到表格时用 text 策略重试，仅保留过门槛的候选。"""
    try:
        candidates = list(page.find_tables(strategy='text').tables)
    except Exception:  # noqa: BLE001
        return []
    kept = []
    for t in candidates:
        try:
            rows = t.extract()
        except Exception:  # noqa: BLE001
            continue
        if _accept_text_table(rows):
            kept.append(t)
    return kept


def _page_to_md(page, default_body_size: float = 11.0):
    try:
        import fitz
    except Exception:  # noqa: BLE001
        return ''
    tables = []
    try:
        tables = list(page.find_tables().tables)
    except Exception:  # noqa: BLE001
        tables = []
    if not tables:
        tables = _retry_text_strategy_tables(page)
    tbl_boxes = []
    for t in tables:
        try:
            tbl_boxes.append(fitz.Rect(t.bbox))
        except Exception:  # noqa: BLE001
            pass

    def in_table(r):
        try:
            area = r.get_area()
            if area <= 0:
                return False
            for tb in tbl_boxes:
                inter = r & tb
                if not inter.is_empty and inter.get_area() > 0.5 * area:
                    return True
        except Exception:  # noqa: BLE001
            return False
        return False

    items = []
    seq = 0
    for t in tables:
        try:
            md = _data_to_md(t.extract())
        except Exception:  # noqa: BLE001
            md = ''
        if md:
            items.append((t.bbox[1], seq, 'table', md))
            seq += 1

    extracted_lines = []
    try:
        d = page.get_text('dict')
        # 统计本页 font size 众数
        sizes = []
        for block in d.get('blocks', []):
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                for sp in line.get('spans', []):
                    sz = sp.get('size', 0)
                    t = (sp.get('text') or '').strip()
                    if sz > 4 and t:
                        sizes.append(round(sz, 1))
        body_size = default_body_size
        if sizes:
            from collections import Counter
            c = Counter(sizes)
            body_size = c.most_common(1)[0][0]

        for block in d.get('blocks', []):
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                bbox = fitz.Rect(line['bbox'])
                if in_table(bbox):
                    continue
                spans = line.get('spans', [])
                if not spans:
                    continue
                line_text = ''.join((sp.get('text') or '') for sp in spans).rstrip()
                if not line_text.strip():
                    continue

                max_size = max(sp.get('size', body_size) for sp in spans)
                is_bold = any(bool(sp.get('flags', 0) & 2) or 'bold' in (sp.get('font') or '').lower() for sp in spans)
                is_mono = any(_is_mono_font(sp.get('font') or '') for sp in spans)

                extracted_lines.append({
                    'y0': bbox.y0,
                    'seq': seq,
                    'text': line_text,
                    'size': max_size,
                    'bold': is_bold,
                    'mono': is_mono,
                    'body_size': body_size
                })
                seq += 1
    except Exception:  # noqa: BLE001
        pass

    idx = 0
    while idx < len(extracted_lines):
        cur = extracted_lines[idx]
        txt = cur['text'].strip()

        # 1. 代码块判断（等宽字体）
        if cur['mono']:
            code_lines = [cur['text']]
            j = idx + 1
            while j < len(extracted_lines) and extracted_lines[j]['mono']:
                code_lines.append(extracted_lines[j]['text'])
                j += 1
            code_content = '\n'.join(code_lines)
            items.append((cur['y0'], cur['seq'], 'code', '```\n' + code_content + '\n```'))
            idx = j
            continue

        # 2. 标题判断
        bs = cur['body_size']
        sz = cur['size']
        is_bold = cur['bold']
        is_short = len(txt) <= 80 and not txt.endswith(('。', '.', '；', ';', '，', ','))

        if sz >= 1.45 * bs and is_short:
            items.append((cur['y0'], cur['seq'], 'heading', '# ' + txt))
            idx += 1
            continue
        elif sz >= 1.25 * bs and is_short:
            items.append((cur['y0'], cur['seq'], 'heading', '## ' + txt))
            idx += 1
            continue
        elif (sz >= 1.12 * bs or (is_bold and sz >= bs)) and is_short:
            items.append((cur['y0'], cur['seq'], 'heading', '### ' + txt))
            idx += 1
            continue

        # 3. 列表判断
        if _BULLET_PREFIX_RE.match(txt):
            clean_item = _BULLET_PREFIX_RE.sub('', txt).strip()
            items.append((cur['y0'], cur['seq'], 'list', '- ' + clean_item))
            idx += 1
            continue
        elif _ORDERED_PREFIX_RE.match(txt):
            items.append((cur['y0'], cur['seq'], 'list', txt))
            idx += 1
            continue

        # 4. 公式判断
        if _looks_like_formula(txt):
            items.append((cur['y0'], cur['seq'], 'formula', '$' + txt + '$'))
            idx += 1
            continue

        # 5. 普通正文段落拼接
        para_text = cur['text']
        j = idx + 1
        while j < len(extracted_lines):
            nxt = extracted_lines[j]
            nxt_txt = nxt['text'].strip()
            if nxt['mono'] or nxt['size'] >= 1.12 * bs or _BULLET_PREFIX_RE.match(nxt_txt) or _ORDERED_PREFIX_RE.match(nxt_txt) or _looks_like_formula(nxt_txt):
                break
            if abs(nxt['y0'] - extracted_lines[j-1]['y0']) > (nxt['size'] * 2.2):
                break
            para_text = _join_lines(para_text, nxt['text'])
            j += 1

        items.append((cur['y0'], cur['seq'], 'text', para_text.strip()))
        idx = j

    items.sort(key=lambda it: (round(it[0], 1), it[1]))
    out = []
    for _y0, _seq, kind, payload in items:
        out.append(payload)
        out.append('')
    return '\n'.join(out).strip()


def pdf2md(path):
    import fitz
    doc = fitz.open(path)
    parts = []
    try:
        all_sizes = []
        for page in doc:
            try:
                d = page.get_text('dict')
                for block in d.get('blocks', []):
                    if block.get('type') != 0:
                        continue
                    for line in block.get('lines', []):
                        for sp in line.get('spans', []):
                            sz = sp.get('size', 0)
                            t = (sp.get('text') or '').strip()
                            if sz > 4 and t:
                                all_sizes.append(round(sz, 1))
            except Exception:
                pass
        global_body_size = 11.0
        if all_sizes:
            from collections import Counter
            global_body_size = Counter(all_sizes).most_common(1)[0][0]

        for page in doc:
            p = _page_to_md(page, default_body_size=global_body_size)
            if p.strip():
                parts.append(p)
    finally:
        doc.close()
    text = '\n\n'.join(parts).strip()
    if text:
        text = _merge_split_tables(text)
    if not text:
        raise ValueError('pdf 未提取到文字内容（可能是扫描件，请用 OCR）')
    return text


def extract_zip_archive(zip_source, base_temp_dir=None):
    """解压 ZIP 归档文件到临时目录，带完整安全管线：
    1. 防 Zip Bomb：解压上限 500MB，单文件/数量/路径深度限制；
    2. 防路径穿越与拒绝符号链接/设备文件；
    3. 支持格式过滤：仅解压 Markdown、文档、文本代码及可转换二进制/图片；
    4. 中文文件名编码自适应（cp437/gbk/utf-8）；
    5. 返回提取路径与跳过统计诊断。
    """
    import zipfile
    import uuid
    import io
    import posixpath
    import time as _time

    MAX_ZIP_SIZE = 100 * 1024 * 1024          # 100 MB 压缩包大小限制
    MAX_TOTAL_UNCOMPRESSED = 500 * 1024 * 1024 # 500 MB 解压总大小限制 (防炸弹)
    MAX_FILE_COUNT = 1000                      # 最多提取 1000 个文件
    MAX_PATH_DEPTH = 12                        # 最深 12 层目录

    SUPPORTED_EXTS = {
        '.md', '.markdown', '.mdown', '.txt', '.text',
        '.json', '.csv', '.tsv', '.yaml', '.yml', '.xml', '.sql',
        '.py', '.js', '.ts', '.html', '.css', '.c', '.cpp', '.h', '.rs', '.go', '.java', '.sh', '.bat', '.ps1',
        '.doc', '.docx', '.ppt', '.pptx', '.xls', '.xlsx', '.pdf', '.epub', '.mobi', '.rtf', '.odt',
        '.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif'
    }

    if not base_temp_dir:
        data_dir = os.environ.get('READMD_DATA_DIR') or os.path.expanduser('~/.readmd')
        base_temp_dir = os.path.join(data_dir, 'temp_zip')

    # Remove abandoned extraction runs from interrupted requests.  The active
    # run is never touched; callers receive paths they can consume normally.
    try:
        os.makedirs(base_temp_dir, exist_ok=True)
        now = _time.time()
        for name in os.listdir(base_temp_dir):
            candidate = os.path.join(base_temp_dir, name)
            if os.path.isdir(candidate) and now - os.path.getmtime(candidate) > 3600:
                shutil.rmtree(candidate, ignore_errors=True)
    except Exception:
        pass

    run_id = str(uuid.uuid4())[:8]
    target_dir = os.path.realpath(os.path.abspath(os.path.join(base_temp_dir, run_id)))
    os.makedirs(target_dir, exist_ok=True)

    extracted_files = []
    skipped_count = 0
    skipped_reasons = {
        'unsupported_format': 0,
        'unsafe_file_type': 0,
        'invalid_path': 0,
        'limit_exceeded': 0
    }

    if isinstance(zip_source, (bytes, bytearray)):
        if len(zip_source) > MAX_ZIP_SIZE:
            raise ValueError('zip_archive_too_large')
        zf_ctx = zipfile.ZipFile(io.BytesIO(zip_source))
    else:
        if os.path.isfile(zip_source) and os.path.getsize(zip_source) > MAX_ZIP_SIZE:
            raise ValueError('zip_archive_too_large')
        zf_ctx = zipfile.ZipFile(zip_source, 'r')

    total_uncompressed = 0
    total_entries = 0

    with zf_ctx as zf:
        infolist = zf.infolist()
        # Bound central-directory metadata before iterating.  This prevents a
        # metadata-only archive from creating an unbounded Python object list.
        if len(infolist) > MAX_FILE_COUNT * 4:
            raise ValueError('zip_entry_count_exceeded')
        total_entries = len(infolist)
        for info in infolist:
            if info.is_dir():
                continue

            # 1. 拒绝符号链接、设备文件与 FIFO
            mode = (info.external_attr >> 16) & 0o170000
            if mode != 0 and mode != 0o100000:  # 0o100000 is regular file
                skipped_count += 1
                skipped_reasons['unsafe_file_type'] += 1
                continue

            # 2. 数量与解压总容量上限
            if len(extracted_files) >= MAX_FILE_COUNT or (total_uncompressed + info.file_size) > MAX_TOTAL_UNCOMPRESSED:
                skipped_count += 1
                skipped_reasons['limit_exceeded'] += 1
                continue

            raw_name = info.filename
            if not (info.flag_bits & 0x800):
                try:
                    raw_bytes = raw_name.encode('cp437')
                    for enc in ('gbk', 'utf-8', 'gb18030', 'big5'):
                        try:
                            raw_name = raw_bytes.decode(enc)
                            break
                        except UnicodeDecodeError:
                            pass
                except Exception:
                    pass

            # Validate the archive name before normalising it.  Stripping a
            # leading slash first would silently turn absolute and UNC names
            # into apparently safe relative paths.
            archive_name = str(raw_name).replace('\\', '/')
            drive, _ = ntpath.splitdrive(archive_name)
            normalized_name = posixpath.normpath(archive_name)
            if (archive_name.startswith('/') or archive_name.startswith('//')
                    or bool(drive) or normalized_name in ('.', '..')
                    or normalized_name.startswith('../')):
                skipped_count += 1
                skipped_reasons['invalid_path'] += 1
                continue
            clean_name = normalized_name

            # 3. 路径深度限制
            parts = [p for p in clean_name.replace('\\', '/').split('/') if p and p != '.']
            if len(parts) > MAX_PATH_DEPTH:
                skipped_count += 1
                skipped_reasons['invalid_path'] += 1
                continue

            # 4. 支持格式过滤
            ext = os.path.splitext(clean_name)[1].lower()
            if ext not in SUPPORTED_EXTS:
                skipped_count += 1
                skipped_reasons['unsupported_format'] += 1
                continue

            dest_file = os.path.realpath(os.path.abspath(os.path.join(target_dir, clean_name)))
            if not dest_file.startswith(target_dir + os.sep) and dest_file != target_dir:
                skipped_count += 1
                skipped_reasons['invalid_path'] += 1
                continue

            os.makedirs(os.path.dirname(dest_file), exist_ok=True)
            with zf.open(info) as src, open(dest_file, 'wb') as dst:
                while True:
                    chunk = src.read(65536)
                    if not chunk:
                        break
                    dst.write(chunk)

            if os.path.isfile(dest_file):
                extracted_files.append(dest_file)
                total_uncompressed += info.file_size

    return {
        'ok': True,
        'paths': extracted_files,
        'skipped': skipped_count,
        'total': total_entries,
        'reasons': skipped_reasons
    }
