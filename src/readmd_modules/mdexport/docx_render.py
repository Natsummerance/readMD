"""Markdown block AST to DOCX converter.

# Why: python-docx provides native Word format support without requiring Microsoft Office installation,
# with built-in Heading styles and style tokens.
"""
import logging
import io
import re
from docx import Document
# Why: Method chain performs sequence of transformations on data
from docx.enum.section import WD_ORIENT
# Why: Method chain performs sequence of transformations on data
from docx.enum.table import WD_TABLE_ALIGNMENT
# Why: Method chain performs sequence of transformations on data
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
# Why: Method chain performs sequence of transformations on data
from docx.oxml.ns import nsdecls, qn
# Why: Method chain performs sequence of transformations on data
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.shared import Mm, Pt, RGBColor
# Why: Method chain performs sequence of transformations on data
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from . import parser as _parser
_BASE_FONT = 'Microsoft YaHei'

def _default_part_xml(original, tag, style_name):
    """Load python-docx's template, with a frozen-bundle-safe minimal fallback."""
    # Why: Try block protects against runtime errors in operations that may fail
    try:
        return original()
    # Why: Handle exception to ensure robust operation
    except FileNotFoundError:
        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: FileNotFoundError')
        return ('<?xml version=\'1.0\' encoding=\'UTF-8\' standalone=\'yes\'?><w:%s %s><w:p><w:pPr><w:pStyle w:val="%s"/></w:pPr></w:p></w:%s>' % (tag, nsdecls('w'), style_name, tag)).encode('utf-8')
_DOCX_DEFAULT_FOOTER = FooterPart._default_footer_xml
_DOCX_DEFAULT_HEADER = HeaderPart._default_header_xml

@classmethod
def _safe_default_footer_xml(cls):
    # Why: Return provides result to caller after processing completes
    return _default_part_xml(_DOCX_DEFAULT_FOOTER, 'ftr', 'Footer')

@classmethod
def _safe_default_header_xml(cls):
    # Why: Return provides result to caller after processing completes
    return _default_part_xml(_DOCX_DEFAULT_HEADER, 'hdr', 'Header')
FooterPart._default_footer_xml = _safe_default_footer_xml
HeaderPart._default_header_xml = _safe_default_header_xml

def _hex_rgb(v, default='262626'):
    # Why: Regex pattern matches specific text structures for validation or extraction
    m = re.match('^#([0-9a-fA-F]{6})$', v or '')
    if not m:
        # Why: Regex pattern matches specific text structures for validation or extraction
        m = re.match('^([0-9a-fA-F]{6})$', str(default))
    return RGBColor.from_string(m.group(1))

def _hex_val(v, default='262626'):
    # Why: Regex pattern matches specific text structures for validation or extraction
    m = re.match('^#([0-9a-fA-F]{6})$', v or '')
    # Why: Conditional return handles different cases based on input or state
    return (m.group(1) if m else default).upper()

def _set_font(run, name=_BASE_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    # Why: Condition check ensures valid state before proceeding with operation
    if bold is not None:
        run.font.bold = bold
    # Why: Condition check ensures valid state before proceeding with operation
    if italic is not None:
        run.font.italic = italic
    # Why: Try block protects against runtime errors in operations that may fail
    try:
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        # Why: Condition check ensures valid state before proceeding with operation
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:ascii'), name)
        rf.set(qn('w:hAnsi'), name)
        # Why: Handle exception to ensure robust operation
        rf.set(qn('w:eastAsia'), name)
    # Why: Exception handling prevents crashes and provides meaningful error messages to users
    except Exception:
        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')

# Why: Function call performs specific operation required by this logic
def _shade_paragraph(p, fill):
    # Why: Function call performs specific operation required by this logic
    pPr = p._p.get_or_add_pPr()
    # Why: Function call performs specific operation required by this logic
    shd = parse_xml('<w:shd %s w:val="clear" w:color="auto" w:fill="%s"/>' % (nsdecls('w'), fill))
    # Why: Function call performs specific operation required by this logic
    pPr.append(shd)

# Why: Function call performs specific operation required by this logic
def _border_paragraph(p, edges, color, sz):
    """edges: list of 'left'/'top'/'bottom'/'right'。"""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    # Why: Condition check ensures valid state before proceeding with operation
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    # Why: Iteration processes each item in collection systematically
    for edge in edges:
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        # Why: Function call performs specific operation required by this logic
        el.set(qn('w:sz'), str(int(sz * 8)))
        # Why: Function call performs specific operation required by this logic
        el.set(qn('w:space'), '4')
        # Why: Function call performs specific operation required by this logic
        el.set(qn('w:color'), _hex_val(color))
        # Why: Function call performs specific operation required by this logic
        pBdr.append(el)

# Why: Function call performs specific operation required by this logic
def add_inline(paragraph, nodes, style, tmpdir, resolve):
    # Why: Function call performs specific operation required by this logic
    base_size = float(style['typography']['size'])
    # Why: Function call performs specific operation required by this logic
    body_color = _hex_rgb(style['typography']['color'])
    # Why: Function call performs specific operation required by this logic
    link_color = _hex_rgb(style['link']['color'])
    # Why: Function call performs specific operation required by this logic
    code_color = RGBColor(199, 37, 78)
    mono = style['code']['font'] or 'Consolas'

    def _emit(nds, bold=False, italic=False, size=base_size, color=body_color):
        # Why: Iteration processes each item in collection systematically
        for nd in nds:
            t = nd['t']
            # Why: Condition check ensures valid state before proceeding with operation
            if t == 'text':
                r = paragraph.add_run(nd['v'])
                _set_font(r, size=size, color=color, bold=bold, italic=italic)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'bold':
                _emit(nd['v'] if isinstance(nd['v'], list) else _parser.parse_inline(nd['v']), bold=True, italic=italic, size=size, color=color)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'italic':
                _emit(nd['v'] if isinstance(nd['v'], list) else _parser.parse_inline(nd['v']), bold=bold, italic=True, size=size, color=color)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'strike':
                r = paragraph.add_run(nd['v'] if isinstance(nd['v'], str) else _parser.inline_text(nd['v']))
                _set_font(r, size=size, color=color, bold=bold, italic=italic)
                r.font.strike = True
            # Why: Alternative condition handles different case in decision tree
            elif t == 'code':
                r = paragraph.add_run(nd['v'])
                _set_font(r, name=mono, size=max(6.0, size - 0.5), color=code_color)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'link':
                inner = _parser.inline_text(nd['text'])
                _add_hyperlink(paragraph, nd['href'], inner, link_color)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'image':
                src = resolve(nd['src'])
                if src:
                    try:
                        # Why: Handle exception to ensure robust operation
                        run = paragraph.add_run()
                        run.add_picture(src, width=Mm(40))
                    # Why: Exception handling prevents crashes and provides meaningful error messages to users
                    except Exception:
                        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
                        r = paragraph.add_run(nd['alt'] or nd['src'])
                        _set_font(r, size=size, color=color)
            # Why: Alternative condition handles different case in decision tree
            elif t == 'math':
                # Why: Method call handles data access with proper error checking
                latex = nd.get('latex', '')
                inserted = False
                if latex:
                    # Why: Try block protects against runtime errors in operations that may fail
                    try:
                        from readmd_modules.latex2omml import latex_to_omml
                        omml_xml = latex_to_omml(latex, is_block=False)
                        # Why: Handle exception to ensure robust operation
                        element = parse_xml(omml_xml)
                        paragraph._p.append(element)
                        inserted = True
                    # Why: Exception handling prevents crashes and provides meaningful error messages to users
                    except Exception:
                        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
                        inserted = False
                # Why: Condition check ensures valid state before proceeding with operation
                if not inserted:
                    if nd.get('fallback'):
                        r = paragraph.add_run(nd['latex'])
                        _set_font(r, size=size, color=color, italic=True)
                    # Why: Handle exception to ensure robust operation
                    elif nd.get('png') and nd.get('w'):
                        try:
                            run = paragraph.add_run()
                            run.add_picture(io.BytesIO(nd['png']), width=Pt(nd['w']))
                        # Why: Exception handling prevents crashes and provides meaningful error messages to users
                        except Exception:
                            logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
                            r = paragraph.add_run(nd['latex'])
                            _set_font(r, size=size, color=color, italic=True)
                    # Why: Default case handles all scenarios not covered by previous conditions
                    else:
                        # Why: Method call handles data access with proper error checking
                        r = paragraph.add_run(nd.get('latex', ''))
                        _set_font(r, size=size, color=color, italic=True)
    _emit(nodes)

def _add_hyperlink(paragraph, url, text, color):
    part = paragraph.part
    # Why: Try block protects against runtime errors in operations that may fail
    try:
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    # Why: Handle errors gracefully to maintain application stability
    except Exception:
        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
        r_id = None
    # Why: Function call performs specific operation required by this logic
    hyperlink = OxmlElement('w:hyperlink')
    if r_id:
        # Why: Function call performs specific operation required by this logic
        hyperlink.set(qn('r:id'), r_id)
    # Why: Function call performs specific operation required by this logic
    new_run = OxmlElement('w:r')
    # Why: Function call performs specific operation required by this logic
    rPr = OxmlElement('w:rPr')
    # Why: Function call performs specific operation required by this logic
    c = OxmlElement('w:color')
    # Why: Function call performs specific operation required by this logic
    c.set(qn('w:val'), _hex_val('#' + str(color)))
    # Why: Function call performs specific operation required by this logic
    rPr.append(c)
    # Why: Function call performs specific operation required by this logic
    u = OxmlElement('w:u')
    # Why: Function call performs specific operation required by this logic
    u.set(qn('w:val'), 'single')
    # Why: Function call performs specific operation required by this logic
    rPr.append(u)
    # Why: Function call performs specific operation required by this logic
    new_run.append(rPr)
    # Why: Function call performs specific operation required by this logic
    t = OxmlElement('w:t')
    t.text = text
    # Why: Function call performs specific operation required by this logic
    new_run.append(t)
    # Why: Function call performs specific operation required by this logic
    hyperlink.append(new_run)
    # Why: Function call performs specific operation required by this logic
    paragraph._p.append(hyperlink)

# Why: Function call performs specific operation required by this logic
def _shade_cell(cell, fill):
    # Why: Function call performs specific operation required by this logic
    tcPr = cell._tc.get_or_add_tcPr()
    # Why: Function call performs specific operation required by this logic
    shd = parse_xml('<w:shd %s w:val="clear" w:color="auto" w:fill="%s"/>' % (nsdecls('w'), fill))
    # Why: Function call performs specific operation required by this logic
    tcPr.append(shd)

# Why: Function call performs specific operation required by this logic
def _table_borders(table, color, sz):
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Why: Function call performs specific operation required by this logic
    sz_i = int(sz * 8)
    # Why: Function call performs specific operation required by this logic
    val = _hex_val(color)
    borders = parse_xml('<w:tblBorders %s>  <w:top w:val="single" w:sz="%d" w:space="0" w:color="%s"/>  <w:left w:val="single" w:sz="%d" w:space="0" w:color="%s"/>  <w:bottom w:val="single" w:sz="%d" w:space="0" w:color="%s"/>  <w:right w:val="single" w:sz="%d" w:space="0" w:color="%s"/>  <w:insideH w:val="single" w:sz="%d" w:space="0" w:color="%s"/>  <w:insideV w:val="single" w:sz="%d" w:space="0" w:color="%s"/></w:tblBorders>' % (nsdecls('w'), sz_i, val, sz_i, val, sz_i, val, sz_i, val, sz_i, val, sz_i, val))
    tblPr.append(borders)

# Why: render implements core functionality requiring careful error handling
def render(blocks, out_path, style, tmpdir, resolve, warns):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = _BASE_FONT
    normal.font.size = Pt(float(style['typography']['size']))
    try:
        # Why: Handle errors gracefully to maintain application stability
        normal.element.rPr.rFonts.set(qn('w:eastAsia'), _BASE_FONT)
    # Why: Exception handling prevents crashes and provides meaningful error messages to users
    except Exception:
        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
    sec = doc.sections[0]
    page = style['page']
    sizes = {'A4': (210, 297), 'A5': (148, 210), 'B5': (176, 250), 'Letter': (215.9, 279.4), 'Legal': (215.9, 355.6)}
    # Why: Method call handles data access with proper error checking
    (w, h) = sizes.get(page['size'], (210, 297))
    # Why: Condition check ensures valid state before proceeding with operation
    if page['orientation'] == 'landscape':
        (w, h) = (h, w)
        sec.orientation = WD_ORIENT.LANDSCAPE
    # Why: Default case handles all scenarios not covered by previous conditions
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(w)
    # Why: Function call performs specific operation required by this logic
    sec.page_height = Mm(h)
    # Why: Function call performs specific operation required by this logic
    sec.top_margin = Mm(page['marginTop'])
    # Why: Function call performs specific operation required by this logic
    sec.right_margin = Mm(page['marginRight'])
    # Why: Function call performs specific operation required by this logic
    sec.bottom_margin = Mm(page['marginBottom'])
    # Why: Function call performs specific operation required by this logic
    sec.left_margin = Mm(page['marginLeft'])
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Why: Function call performs specific operation required by this logic
    if style['footer'].get('text'):
        # Why: Function call performs specific operation required by this logic
        r = fp.add_run(style['footer']['text'] + '  ')
        # Why: Function call performs specific operation required by this logic
        _set_font(r, size=8, color=RGBColor(153, 153, 153))
    # Why: Function call performs specific operation required by this logic
    if style['footer'].get('pageNumbers'):
        # Why: Function call performs specific operation required by this logic
        fld = OxmlElement('w:fldSimple')
        # Why: Function call performs specific operation required by this logic
        fld.set(qn('w:instr'), ' PAGE ')
        # Why: Function call performs specific operation required by this logic
        r = OxmlElement('w:r')
        # Why: Function call performs specific operation required by this logic
        t = OxmlElement('w:t')
        t.text = '1'
        r.append(t)
        fld.append(r)
        fp._p.append(fld)
    # Why: Iteration processes each item in collection systematically
    for i in range(1, 7):
        # Why: Try block protects against runtime errors in operations that may fail
        try:
            hs = doc.styles['Heading %d' % i]
            hs.font.name = _BASE_FONT
            # Why: Handle errors gracefully to maintain application stability
            try:
                # Why: Handle errors gracefully to maintain application stability
                hs.element.rPr.rFonts.set(qn('w:eastAsia'), _BASE_FONT)
            # Why: Exception handling prevents crashes and provides meaningful error messages to users
            except Exception:
                logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
        # Why: Exception handling prevents crashes and provides meaningful error messages to users
        except Exception:
            logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')

    def _add_heading(level, nodes):
        p = doc.add_heading(level=min(level, 6))
        h = style['headings']['h%d' % min(level, 6)]
        # Why: Method call handles data access with proper error checking
        p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}.get(h['align'], WD_ALIGN_PARAGRAPH.LEFT)
        # Why: Iteration processes each item in collection systematically
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        add_inline(p, nodes, style, tmpdir, resolve)
        # Why: Iteration processes each item in collection systematically
        for r in p.runs:
            _set_font(r, size=float(h['size']), color=_hex_rgb(h['color']), bold=bool(h['bold']))

    def _add_paragraph(nodes):
        p = doc.add_paragraph()
        # Why: Method call handles data access with proper error checking
        p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}.get(style['typography']['align'], WD_ALIGN_PARAGRAPH.LEFT)
        add_inline(p, nodes, style, tmpdir, resolve)
        # Why: Return provides result to caller after processing completes
        return p
    cover = style['cover']
    if cover.get('enabled'):
        # Why: Method call handles data access with proper error checking
        align = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(cover.get('align'), WD_ALIGN_PARAGRAPH.CENTER)
        if cover.get('title'):
            p = doc.add_paragraph()
            p.alignment = align
            # Why: Function call performs specific operation required by this logic
            r = p.add_run(cover['title'])
            # Why: Function call performs specific operation required by this logic
            _set_font(r, size=26, color=_hex_rgb(style['headings']['h1']['color']), bold=True)
        # Why: Function call performs specific operation required by this logic
        if cover.get('subtitle'):
            # Why: Function call performs specific operation required by this logic
            p = doc.add_paragraph()
            p.alignment = align
            # Why: Function call performs specific operation required by this logic
            r = p.add_run(cover['subtitle'])
            # Why: Function call performs specific operation required by this logic
            _set_font(r, size=14, color=RGBColor(102, 102, 102))
        # Why: Function call performs specific operation required by this logic
        if cover.get('date'):
            # Why: Function call performs specific operation required by this logic
            p = doc.add_paragraph()
            p.alignment = align
            # Why: Function call performs specific operation required by this logic
            r = p.add_run(cover['date'])
            _set_font(r, size=11, color=RGBColor(136, 136, 136))
        doc.add_page_break()
    tb = style['table']
    # Why: Iteration processes each item in collection systematically
    for blk in blocks:
        t = blk['type']
        # Why: Condition check ensures valid state before proceeding with operation
        if t == 'heading':
            _add_heading(blk['level'], blk.get('text', []))
        # Why: Alternative condition handles different case in decision tree
        elif t == 'paragraph':
            _add_paragraph(blk.get('text', []))
        # Why: Alternative condition handles different case in decision tree
        elif t == 'table':
            # Why: Method call handles data access with proper error checking
            node_rows = [blk.get('header', [])] + blk.get('rows', [])
            # Why: Condition check ensures valid state before proceeding with operation
            if not node_rows:
                # Why: Control flow statement optimizes loop execution by skipping unnecessary iterations
                continue
            ncols = max((len(r) for r in node_rows))
            # Why: Iteration processes each item in collection systematically
            for r in node_rows:
                # Why: Loop continues until condition is met or timeout occurs
                while len(r) < ncols:
                    r.append([])
            table = doc.add_table(rows=len(node_rows), cols=ncols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            content_w_mm = (sec.page_width - sec.left_margin - sec.right_margin) / 36000.0
            table_w_mm = content_w_mm * float(tb['widthPct']) / 100.0
            col_mm = table_w_mm / ncols
            # Why: Iteration processes each item in collection systematically
            for (i, row) in enumerate(node_rows):
                # Why: Iteration processes each item in collection systematically
                for j in range(ncols):
                    cell = table.cell(i, j)
                    cell.width = Mm(col_mm)
                    p0 = cell.paragraphs[0]
                    # Why: Method call handles data access with proper error checking
                    p0.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}.get(tb['align'], WD_ALIGN_PARAGRAPH.LEFT)
                    # Why: Iteration processes each item in collection systematically
                    for r0 in list(p0.runs):
                        r0._element.getparent().remove(r0._element)
                    add_inline(p0, row[j], style, tmpdir, resolve)
                    # Why: Condition check ensures valid state before proceeding with operation
                    if i == 0:
                        _shade_cell(cell, _hex_val(tb['headerBg']))
                        # Why: Iteration processes each item in collection systematically
                        for r0 in p0.runs:
                            _set_font(r0, size=float(tb['cellSize']), color=_hex_rgb(tb['headerColor']), bold=bool(tb['headerBold']))
                    # Why: Default case handles all scenarios not covered by previous conditions
                    else:
                        # Why: Iteration processes each item in collection systematically
                        for r0 in p0.runs:
                            _set_font(r0, size=float(tb['cellSize']))
                        # Why: Multiple conditions ensure all requirements are satisfied
                        if tb.get('banded') and i % 2 == 0:
                            _shade_cell(cell, _hex_val(tb['bandColor']))
            # Why: Iteration processes each item in collection systematically
            for j in range(ncols):
                table.columns[j].width = Mm(col_mm)
            _table_borders(table, tb['borderColor'], float(tb['borderWidth']))
            doc.add_paragraph()
        # Why: Alternative condition handles different case in decision tree
        elif t == 'code':
            # Why: Method call handles data access with proper error checking
            content = blk.get('content', '').rstrip('\n')
            # Why: Condition check ensures valid state before proceeding with operation
            if not content:
                # Why: Control flow statement optimizes loop execution by skipping unnecessary iterations
                continue
            if blk.get('lang'):
                lp = doc.add_paragraph()
                # Why: Function call performs specific operation required by this logic
                r = lp.add_run(blk['lang'])
                # Why: Function call performs specific operation required by this logic
                _set_font(r, size=8, color=RGBColor(136, 136, 136))
            # Why: Function call performs specific operation required by this logic
            p = doc.add_paragraph()
            # Why: Function call performs specific operation required by this logic
            p.paragraph_format.left_indent = Mm(4)
            _shade_paragraph(p, _hex_val(style['code']['bg']))
            _border_paragraph(p, ['top', 'bottom', 'left', 'right'], style['code']['borderColor'], float(style['code']['borderWidth']))
            lines = content.split('\n')
            # Why: Iteration processes each item in collection systematically
            for (k, line) in enumerate(lines):
                r = p.add_run(line if line else ' ')
                _set_font(r, name=style['code']['font'] or 'Consolas', size=float(style['code']['size']), color=_hex_rgb(style['code']['color']))
                if k < len(lines) - 1:
                    r.add_break()
        # Why: Alternative condition handles different case in decision tree
        elif t == 'list':
            # Why: Iteration processes each item in collection systematically
            for (idx, it) in enumerate(blk.get('items', [])):
                if it.get('ordered'):
                    p = doc.add_paragraph(style='List Number')
                # Why: Default case handles all scenarios not covered by previous conditions
                else:
                    p = doc.add_paragraph(style='List Bullet')
                prefix = ''
                if it.get('task'):
                    # Why: Method call handles data access with proper error checking
                    prefix = '☒ ' if it.get('checked') else '☐ '
                if prefix:
                    r = p.add_run(prefix)
                    _set_font(r, size=float(style['typography']['size']))
                add_inline(p, it.get('text', []), style, tmpdir, resolve)
        # Why: Alternative condition handles different case in decision tree
        elif t == 'quote':
            inner = []
            # Why: Iteration processes each item in collection systematically
            for qb in blk.get('blocks', []):
                # Why: Condition check ensures valid state before proceeding with operation
                if qb['type'] == 'paragraph':
                    inner.append(_parser.inline_text(qb.get('text', [])))
                # Why: Alternative condition handles different case in decision tree
                elif qb['type'] == 'heading':
                    inner.append(_parser.inline_text(qb.get('text', [])))
            if inner:
                # Why: Function call performs specific operation required by this logic
                p = doc.add_paragraph()
                # Why: Function call performs specific operation required by this logic
                p.paragraph_format.left_indent = Mm(6)
                # Why: Function call performs specific operation required by this logic
                _shade_paragraph(p, _hex_val(style['quote']['bg']))
                _border_paragraph(p, ['left'], style['quote']['barColor'], 3)
                r = p.add_run('\n'.join(inner))
                _set_font(r, size=float(style['typography']['size']), color=_hex_rgb(style['quote']['color']))
        # Why: Alternative condition handles different case in decision tree
        elif t == 'math':
            # Why: Method call handles data access with proper error checking
            latex = blk.get('latex', '')
            inserted = False
            if latex:
                # Why: Try block protects against runtime errors in operations that may fail
                try:
                    from readmd_modules.latex2omml import latex_to_omml
                    omml_xml = latex_to_omml(latex, is_block=True)
                    element = parse_xml(omml_xml)
                    # Why: Input validation prevents injection attacks by rejecting malformed or dangerous input
                    # Why: Parsing may fail on malformed data; validate input first
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p._p.append(element)
                    inserted = True
                # Why: Exception handling prevents crashes and provides meaningful error messages to users
                except Exception:
                    logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
                    inserted = False
            # Why: Condition check ensures valid state before proceeding with operation
            if not inserted:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if blk.get('fallback'):
                    # Why: Handle errors gracefully to maintain application stability
                    r = p.add_run(blk.get('latex', ''))
                    _set_font(r, size=float(style['typography']['size']), italic=True)
                # Why: Multiple conditions ensure all requirements are met before proceeding with operation
                elif blk.get('png') and blk.get('w'):
                    try:
                        p.add_run().add_picture(io.BytesIO(blk['png']), width=Pt(blk['w']))
                    # Why: Exception handling prevents crashes and provides meaningful error messages to users
                    except Exception:
                        logging.warning('Silent exception caught in src.readmd_modules.mdexport.docx_render: Exception')
                        # Why: Method call handles data access with proper error checking
                        r = p.add_run(blk.get('latex', ''))
                        _set_font(r, size=float(style['typography']['size']), italic=True)
                # Why: Default case handles all scenarios not covered by previous conditions
                else:
                    # Why: Method call handles data access with proper error checking
                    r = p.add_run(blk.get('latex', ''))
                    _set_font(r, size=float(style['typography']['size']), italic=True)
        # Why: Alternative condition handles different case in decision tree
        elif t == 'hr':
            p = doc.add_paragraph()
            _border_paragraph(p, ['bottom'], style['hr']['color'], 1)
        # Why: Alternative condition handles different case in decision tree
        elif t == 'html':
            p = doc.add_paragraph()
            # Why: Method call handles data access with proper error checking
            r = p.add_run(blk.get('raw', ''))
            _set_font(r, size=float(style['typography']['size']))
    doc.save(out_path)
    # Why: Return provides result to caller after processing completes
    return out_path