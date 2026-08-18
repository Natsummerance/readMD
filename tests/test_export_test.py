# -*- coding: utf-8 -*-
"""ReadMD v2.1.0 导出模块测试。

运行：python test_export_test.py   （退出码 0 = 全部通过）
"""

import os
import shutil
import sys
import tempfile
import unittest
from unittest import mock

ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(ROOT, '..'))

from src.readmd_modules.mdexport import parser as P
from src.readmd_modules.mdexport import styles as S
from src.readmd_modules.mdexport import formula as F
from src.readmd_modules.mdexport import docx_render as DOCX
from src.readmd_modules.mdexport import pdf_render as PDF
import src.readmd_modules.mdexport as E
import sys, os; sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..")); import readmd

SAMPLE = '''# 标题一

正文段落 **加粗** *斜体* `行内码` 和 ~~删除~~ [链接](https://example.com) 与图片 ![图](img/a.png)。

## 表格测试

| 名称 | 数量 | 说明 |
| :--- | ---: | :---: |
| 苹果 | 12 | 水果 |
| 香蕉 | 8 | 水果 |

> 引用内容
> 第二行

- 项目甲
- 项目乙
- [x] 已完成任务
- [ ] 未完成任务

1. 第一
2. 第二

```python
def hello():
    return 1
```

公式 $a^2 + b^2 = c^2$ 与展示：

$$\\int_0^1 x^2 dx = \\frac{1}{3}$$

---

尾部段落。
'''


class TestParser(unittest.TestCase):
    def test_blocks(self):
        b = P.parse(SAMPLE)
        types = [x['type'] for x in b]
        self.assertIn('heading', types)
        self.assertIn('table', types)
        self.assertIn('quote', types)
        self.assertIn('list', types)
        self.assertIn('code', types)
        self.assertIn('math', types)
        self.assertIn('hr', types)

    def test_heading_level(self):
        b = P.parse('## 二级\n\n### 三级')
        self.assertEqual(b[0]['type'], 'heading')
        self.assertEqual(b[0]['level'], 2)
        self.assertEqual(b[1]['level'], 3)

    def test_table_align(self):
        b = P.parse(SAMPLE)
        t = [x for x in b if x['type'] == 'table'][0]
        self.assertEqual(len(t['header']), 3)
        self.assertEqual(len(t['rows']), 2)
        self.assertEqual(t['aligns'], ['left', 'right', 'center'])

    def test_code_fence(self):
        b = P.parse(SAMPLE)
        c = [x for x in b if x['type'] == 'code'][0]
        self.assertEqual(c['lang'], 'python')
        self.assertIn('return 1', c['content'])

    def test_list_task(self):
        b = P.parse(SAMPLE)
        lst = [x for x in b if x['type'] == 'list'][0]
        self.assertTrue(any(i.get('task') and i.get('checked') for i in lst['items']))
        self.assertTrue(any(i.get('task') and not i.get('checked') for i in lst['items']))

    def test_inline_nodes(self):
        b = P.parse('**粗** *斜* `码` ~~删~~ [链](https://a.b) $x$')
        p = b[0]['text']
        kinds = [n['t'] for n in p]
        self.assertIn('bold', kinds)
        self.assertIn('italic', kinds)
        self.assertIn('code', kinds)
        self.assertIn('strike', kinds)
        self.assertIn('link', kinds)
        self.assertIn('math', kinds)

    def test_math_not_in_code(self):
        b = P.parse('```\n$a+b$\n```\n\n正文 $c+d$')
        code = [x for x in b if x['type'] == 'code'][0]
        self.assertNotIn('math', code)
        para = [x for x in b if x['type'] == 'paragraph'][0]
        self.assertTrue(any(n['t'] == 'math' for n in para['text']))

    def test_display_math(self):
        b = P.parse('$$\\frac{a}{b}$$')
        m = [x for x in b if x['type'] == 'math'][0]
        self.assertTrue(m['display'])
        self.assertIn('frac', m['latex'])

    def test_inline_text(self):
        nodes = P.parse_inline('**加粗** 和 [链接](x)')
        text = P.inline_text(nodes)
        self.assertIn('加粗', text)
        self.assertIn('链接', text)


class TestStyles(unittest.TestCase):
    def test_sanitize_defaults(self):
        s = S.sanitize({})
        self.assertEqual(s['page']['size'], 'A4')
        self.assertEqual(s['headings']['h1']['size'], 20)
        self.assertEqual(s['htmlTheme'], 'light')

    def test_sanitize_bad_values(self):
        s = S.sanitize({'page': {'size': 'XXL', 'marginTop': 999},
                        'headings': {'h1': {'size': -5, 'color': 'red'}},
                        'htmlTheme': 'neon'})
        self.assertEqual(s['page']['size'], 'A4')
        self.assertEqual(s['page']['marginTop'], 60)  # 上限
        self.assertEqual(s['headings']['h1']['size'], 8)  # 下限
        self.assertEqual(s['headings']['h1']['color'], '#1a1a1a')
        self.assertEqual(s['htmlTheme'], 'light')

    def test_merge_overrides(self):
        s = S.sanitize({'typography': {'size': 12}, 'table': {'banded': False}})
        self.assertEqual(s['typography']['size'], 12)
        self.assertFalse(s['table']['banded'])
        self.assertEqual(s['table']['headerBg'], '#3b6ef5')

    def test_presets_exist(self):
        for name in ('minimal', 'classic', 'business'):
            s = S.preset_style(name)
            self.assertEqual(s['headings']['h1']['size'] > 0, True)
            self.assertEqual(s['htmlTheme'], 'light')


class TestFormula(unittest.TestCase):
    def test_render_png(self):
        data = F.render_latex(r'\frac{a}{b}+\sqrt{x}')
        self.assertTrue(data and data[:4] == b'\x89PNG')
        size = F.png_size(data)
        self.assertTrue(size and size[0] > 0 and size[1] > 0)

    def test_render_cjk(self):
        data = F.render_latex(r'\frac{常数}{n}')
        self.assertTrue(data and data[:4] == b'\x89PNG')

    def test_render_fail_fallback(self):
        self.assertIsNone(F.render_latex(r'\begin{cases}'))


class TestImageResolver(unittest.TestCase):
    def test_resolve(self):
        with tempfile.TemporaryDirectory() as td:
            with open(os.path.join(td, 'a.png'), 'wb') as f:
                f.write(b'x')
            warns = []
            r = E.ImageResolver(td, warns)
            self.assertTrue(r.resolve('a.png'))
            self.assertIsNone(r.resolve('missing.png'))
            self.assertIsNone(r.resolve('http://x/y.png'))
            self.assertTrue(len(warns) >= 1)


class TestPdfFonts(unittest.TestCase):
    def test_missing_system_font_keeps_registered_fallback(self):
        previous = PDF._registered_font_name
        try:
            PDF._registered_font_name = None
            with mock.patch.object(PDF, '_first_existing', return_value='/missing/font.ttf'):
                first = PDF.register_fonts()
                second = PDF.register_fonts()
            self.assertEqual(first, second)
            self.assertTrue(PDF._font_ready(second))
        finally:
            PDF._registered_font_name = previous


class TestDocxTemplates(unittest.TestCase):
    def test_missing_frozen_template_uses_minimal_part(self):
        from docx.oxml import parse_xml

        def missing():
            raise FileNotFoundError('frozen template')

        xml = DOCX._default_part_xml(missing, 'ftr', 'Footer')
        self.assertTrue(parse_xml(xml).tag.endswith('ftr'))


class TestExportSmoke(unittest.TestCase):
    def _sample_with_img(self, td):
        with open(os.path.join(td, 'a.png'), 'wb') as f:
            f.write(b'\x89PNG\r\n\x1a\n' + b'0' * 64)
        return SAMPLE.replace('![图](img/a.png)', '![图](a.png)')

    def test_pdf(self):
        import fitz
        with tempfile.TemporaryDirectory() as td:
            md = self._sample_with_img(td)
            out = os.path.join(td, 'out.pdf')
            r = E.export('pdf', md, td, out, options={'meta': {'title': 'T'}})
            self.assertTrue(r['ok'], r)
            self.assertGreater(r['size'], 0)
            doc = fitz.open(out)
            self.assertGreaterEqual(doc.page_count, 1)
            self.assertEqual(doc.metadata.get('title'), 'T')
            doc.close()

    def test_docx(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as td:
            md = self._sample_with_img(td)
            out = os.path.join(td, 'out.docx')
            r = E.export('docx', md, td, out)
            self.assertTrue(r['ok'], r)
            d = Document(out)
            self.assertTrue(any('标题一' in p.text for p in d.paragraphs))
            self.assertGreaterEqual(len(d.tables), 1)

    def test_html(self):
        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, 'out.html')
            r = E.export('html', SAMPLE, td, out, options={'htmlTheme': 'dark'})
            self.assertTrue(r['ok'], r)
            h = open(out, encoding='utf-8').read()
            self.assertIn('<!DOCTYPE html>', h)
            self.assertIn('marked.parse', h)
            self.assertIn('MathJax', h)
            self.assertIn('标题一', h)
            self.assertIn('--bg', h)

    def test_bad_format(self):
        r = E.export('txt', SAMPLE, '', 'x.txt')
        self.assertFalse(r['ok'])

    def test_missing_image_warn(self):
        with tempfile.TemporaryDirectory() as td:
            md = SAMPLE  # 引用 img/a.png 不存在
            out = os.path.join(td, 'out.pdf')
            r = E.export('pdf', md, td, out)
            self.assertTrue(r['ok'])
            self.assertTrue(any('图片' in w for w in r['warns']))


class TestExportBridge(unittest.TestCase):
    """Catch pywebview SAVE_DIALOG result-shape regressions."""

    class _Window(object):
        def __init__(self, targets):
            self.targets = iter(targets)

        def create_file_dialog(self, *args, **kwargs):
            # pywebview 6.x WinForms returns a one-item tuple for SAVE_DIALOG.
            return (next(self.targets),)

    def test_windows_save_dialog_tuple_exports_all_formats(self):
        with tempfile.TemporaryDirectory() as td:
            targets = [os.path.join(td, 'bridge.' + fmt)
                       for fmt in ('pdf', 'docx', 'html')]
            api = readmd.Api()
            api._window = self._Window(targets)
            for fmt, target in zip(('pdf', 'docx', 'html'), targets):
                result = api.export_doc(fmt, {
                    'content': '# Bridge\n\nExport path contract.',
                    'baseDir': td,
                    'suggestedName': 'bridge',
                    'options': {},
                })
                self.assertTrue(result.get('ok'), (fmt, result))
                self.assertEqual(result.get('path'), target)
                self.assertTrue(os.path.isfile(target))

    def test_dialog_path_normalization_rejects_multiple_targets(self):
        with self.assertRaises(ValueError):
            readmd.normalize_dialog_path(('a.pdf', 'b.pdf'), '.pdf')
        self.assertTrue(readmd.normalize_dialog_path('report', '.pdf').endswith('report.pdf'))

    def test_failed_export_keeps_existing_destination(self):
        from src.readmd_modules.mdexport import html_render

        with tempfile.TemporaryDirectory() as td:
            target = os.path.join(td, 'kept.html')
            with open(target, 'w', encoding='utf-8') as handle:
                handle.write('original')

            def partial_then_fail(content, out_path, *args, **kwargs):
                with open(out_path, 'w', encoding='utf-8') as handle:
                    handle.write('partial')
                raise RuntimeError('renderer stopped')

            with mock.patch.object(html_render, 'render', partial_then_fail):
                result = E.export('html', '# Test', td, target)

            self.assertFalse(result.get('ok'))
            self.assertEqual(result.get('stage'), 'render')
            with open(target, encoding='utf-8') as handle:
                self.assertEqual(handle.read(), 'original')
            self.assertFalse(any('.readmd-' in name for name in os.listdir(td)))


def main():
    suite = unittest.defaultTestLoader.loadTestsFromModule(sys.modules[__name__])
    result = unittest.TextTestRunner(verbosity=1).run(suite)
    sys.exit(0 if result.wasSuccessful() else 1)


if __name__ == '__main__':
    main()
