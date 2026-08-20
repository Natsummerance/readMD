# -*- coding: utf-8 -*-
"""OMML <-> LaTeX 双向公式转换与 DOCX 原生数学对象自动化测试。"""

import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

from src.readmd_modules.latex2omml import latex_to_omml
from src.readmd_modules.convert import _omml_to_latex, docx2md


class TestOmmlDocxConversion(unittest.TestCase):
    """OMML <-> LaTeX 公式与 DOCX 互转测试。"""

    def test_latex_to_omml_basic_structures(self):
        """测试基本公式、分数、上下标、根式转 OMML。"""
        # 1. 分数
        xml = latex_to_omml(r'\frac{a + b}{c^2}')
        self.assertIn('<m:f>', xml)
        self.assertIn('<m:num>', xml)
        self.assertIn('<m:den>', xml)
        self.assertIn('<m:sSup>', xml)

        # 2. 根式
        xml_rad = latex_to_omml(r'\sqrt[3]{x^2 + 1}')
        self.assertIn('<m:rad>', xml_rad)
        self.assertIn('<m:deg>', xml_rad)

        # 3. 复合上下标
        xml_subsup = latex_to_omml(r'x_i^2')
        self.assertTrue('<m:sSubSup>' in xml_subsup or ('<m:sSub>' in xml_subsup and '<m:sSup>' in xml_subsup))

    def test_latex_to_omml_matrices_and_delimiters(self):
        """测试矩阵、定界符与方程组转 OMML。"""
        # 1. 括号矩阵 pmatrix
        xml_mat = latex_to_omml(r'\begin{pmatrix} 1 & 2 \\ 3 & 4 \end{pmatrix}', is_block=True)
        self.assertIn('<m:oMathPara', xml_mat)
        self.assertIn('<m:m>', xml_mat)
        self.assertIn('<m:mr>', xml_mat)

        # 2. 定界符
        xml_delim = latex_to_omml(r'\left( \frac{1}{x} \right)')
        self.assertIn('<m:d>', xml_delim)
        self.assertIn('<m:begChr', xml_delim)

    def test_latex_to_omml_symbols_and_greek(self):
        """测试希腊字母与数学符号转 OMML。"""
        xml_sym = latex_to_omml(r'\alpha + \beta \le \gamma \times \pi \to \infty')
        self.assertIn('α', xml_sym)
        self.assertIn('β', xml_sym)
        self.assertIn('≤', xml_sym)
        self.assertIn('×', xml_sym)
        self.assertIn('π', xml_sym)
        self.assertIn('→', xml_sym)
        self.assertIn('∞', xml_sym)

    def test_docx_export_and_import_roundtrip(self):
        """测试 MD 导出为 DOCX（包含原生 OMML）以及 DOCX 重新转为 MD（完美提取公式）。"""
        try:
            from docx import Document
            from docx.oxml import parse_xml
        except ImportError:
            self.skipTest("python-docx 依赖未安装")

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'test_math.docx')
            doc = Document()

            # 1. 创建包含原生 OMML 行内公式的段落
            p1 = doc.add_paragraph()
            p1.add_run('根据牛顿第二定律，物体受力满足关系 ')
            omml_inline = latex_to_omml(r'F = m \cdot a', is_block=False)
            p1._p.append(parse_xml(omml_inline))
            p1.add_run('，其中 m 为质量。')

            # 2. 创建包含原生 OMML 独立公式块的段落
            p2 = doc.add_paragraph()
            omml_block = latex_to_omml(r'\int_0^\infty e^{-x} dx = 1', is_block=True)
            p2._p.append(parse_xml(omml_block))

            # 3. 创建包含矩阵公式的段落
            p3 = doc.add_paragraph()
            p3.add_run('线性变换矩阵为：')
            omml_mat = latex_to_omml(r'\begin{pmatrix} a & b \\ c & d \end{pmatrix}', is_block=True)
            p3._p.append(parse_xml(omml_mat))

            doc.save(docx_path)
            self.assertTrue(os.path.exists(docx_path))

            # 4. 执行 docx2md 转换
            md_text = docx2md(docx_path)

            # 5. 验证提取结果
            self.assertIn('根据牛顿第二定律，物体受力满足关系', md_text)
            self.assertTrue('$F = m' in md_text or '$F=m' in md_text)
            self.assertIn('其中 m 为质量', md_text)

            # 独立积分公式
            self.assertIn(r'\int', md_text)
            self.assertIn('e', md_text)

            # 矩阵公式
            self.assertTrue(r'\begin{pmatrix}' in md_text or r'\begin{matrix}' in md_text)
            self.assertIn('a & b', md_text)
            self.assertIn('c & d', md_text)


if __name__ == '__main__':
    unittest.main()
