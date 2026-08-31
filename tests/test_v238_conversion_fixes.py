# -*- coding: utf-8 -*-
"""V2.3.8 转换链路修复回归：合成 fixture，不含任何真实文档内容。"""
import os
import sys
import tempfile
import unittest
from unittest.mock import patch

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

from src.readmd_modules import convert as CV


def _pipe_rows(md):
    out = []
    for line in md.splitlines():
        s = line.strip()
        if s.startswith('|') and s.endswith('|'):
            cells = [c.strip() for c in s[1:-1].split('|')]
            if cells and all(c == '---' for c in cells):
                continue
            out.append(cells)
    return out


class TestDocxMergedCells(unittest.TestCase):
    def _make_docx(self, path):
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.table import _Cell

        d = Document()
        t = d.add_table(rows=3, cols=3)

        def tc_of(row, col):
            return t.rows[row]._tr.findall(qn('w:tc'))[col]

        def set_text(row, col, text):
            _Cell(tc_of(row, col), t).paragraphs[0].text = text

        # 行 0：tc0 横向合并 2 列，删除第 3 个 tc（直接操作 tr，绕开 row.cells 的展开语义）
        set_text(0, 0, '表头跨列')
        set_text(0, 1, '评语')
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), '2')
        tc_of(0, 0).get_or_add_tcPr().append(gs)
        t.rows[0]._tr.remove(tc_of(0, 2))

        # 行 1：tc1 纵向合并起点
        set_text(1, 0, '课程')
        set_text(1, 1, '优')
        vm = OxmlElement('w:vMerge')
        vm.set(qn('w:val'), 'restart')
        tc_of(1, 1).get_or_add_tcPr().append(vm)
        set_text(1, 2, '备注')

        # 行 2：tc1 纵向合并延续（带占位文字，验证应被起点文本覆盖）
        set_text(2, 0, '实践')
        set_text(2, 1, '占位续格')
        tc_of(2, 1).get_or_add_tcPr().append(OxmlElement('w:vMerge'))
        set_text(2, 2, '良')

        d.save(path)

    def test_gridspan_and_vmerge_alignment(self):
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'merged.docx')
            self._make_docx(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            rows = _pipe_rows(text)
            self.assertEqual(rows[0], ['表头跨列', '', '评语'])
            self.assertEqual(rows[1], ['课程', '优', '备注'])
            self.assertEqual(rows[2], ['实践', '优', '良'])


class TestDocxFormTableHeuristic(unittest.TestCase):
    def _make_docx(self, path):
        from docx import Document

        d = Document()
        d.add_heading('示例记录表', level=1)
        t1 = d.add_table(rows=2, cols=2)
        for row, (k, v) in zip(t1.rows, [('项目', '示例内容'), ('日期', '示例日期')]):
            row.cells[0].paragraphs[0].text = k
            row.cells[1].paragraphs[0].text = v
        t2 = d.add_table(rows=2, cols=2)
        long_val = 'A' * 40
        for row, (k, v) in zip(t2.rows, [('编号', long_val), ('备注', '示例备注')]):
            row.cells[0].paragraphs[0].text = k
            row.cells[1].paragraphs[0].text = v
        d.save(path)

    def test_form_table_converted_to_keyvalue(self):
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'form.docx')
            self._make_docx(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('- **项目**: 示例内容', text)
            self.assertIn('- **日期**: 示例日期', text)
            self.assertNotIn('| 项目', text)

    def test_long_cell_table_kept_as_table(self):
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'form.docx')
            self._make_docx(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertTrue(any(line.strip().startswith('| 编号') for line in text.splitlines()))

    def test_form_tables_false_keeps_table(self):
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'form.docx')
            self._make_docx(p)
            text, engine, err = CV.convert_verbose(p, form_tables=False)
            self.assertEqual(engine, 'docx', err)
            self.assertNotIn('- **项目**', text)
            self.assertIn('| 项目', text)


class TestDocxNestedLists(unittest.TestCase):
    def _set_numpr(self, p, ilvl):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        ppr = p._p.get_or_add_pPr()
        numpr = OxmlElement('w:numPr')
        ilvl_el = OxmlElement('w:ilvl')
        ilvl_el.set(qn('w:val'), str(ilvl))
        numid = OxmlElement('w:numId')
        numid.set(qn('w:val'), '1')
        numpr.append(ilvl_el)
        numpr.append(numid)
        ppr.append(numpr)

    def test_ilvl_indents_bullets(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'nested.docx')
            d = Document()
            self._set_numpr(d.add_paragraph('顶层'), 0)
            self._set_numpr(d.add_paragraph('二级'), 1)
            self._set_numpr(d.add_paragraph('三级'), 2)
            d.save(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('- 顶层', text)
            self.assertIn('  - 二级', text)
            self.assertIn('    - 三级', text)

    def test_list_number_style_emits_ordered(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'ordered.docx')
            d = Document()
            d.add_paragraph('第一项', style='List Number')
            d.add_paragraph('第二项', style='List Number')
            d.save(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('1. 第一项', text)
            self.assertIn('2. 第二项', text)

    def test_ordered_counter_resets_after_paragraph(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'ordered.docx')
            d = Document()
            d.add_paragraph('第一组甲', style='List Number')
            d.add_paragraph('第一组乙', style='List Number')
            d.add_paragraph('普通段落')
            d.add_paragraph('第二组甲', style='List Number')
            d.save(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('2. 第一组乙', text)
            self.assertIn('1. 第二组甲', text)
            self.assertNotIn('3. 第二组甲', text)


    def test_ordered_counter_resets_after_heading(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'ordered.docx')
            d = Document()
            d.add_paragraph('第一组甲', style='List Number')
            d.add_heading('小节标题', level=1)
            d.add_paragraph('第二组甲', style='List Number')
            d.save(p)
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'docx', err)
            self.assertIn('1. 第二组甲', text)
            self.assertNotIn('2. 第二组甲', text)


OLE2_MAGIC = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'


class TestDocConversion(unittest.TestCase):
    """V2.3.8 修复 #4：.doc 分支（magic bytes → Word COM → soffice → 稳定错误码）。"""

    def _write(self, td, name, data):
        p = os.path.join(td, name)
        with open(p, 'wb') as f:
            f.write(data)
        return p

    def test_doc2md_rejects_non_ole2_garbage(self):
        with tempfile.TemporaryDirectory() as td:
            p = self._write(td, 'garbage.doc', b'not a word document at all')
            text, err = CV.doc2md(p)
            self.assertEqual(text, '')
            self.assertIn('doc-not-ole2', err)

    def test_convert_verbose_doc_garbage_stable_error(self):
        with tempfile.TemporaryDirectory() as td:
            p = self._write(td, 'garbage.doc', b'\x00\x01\x02junkjunk')
            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(text, '')
            self.assertIn('doc-not-ole2', err)

    def test_doc2md_no_engine_stable_error(self):
        orig_word, orig_soffice = CV._doc2docx_word_com, CV._doc2docx_soffice
        try:
            CV._doc2docx_word_com = lambda src, out_dir: None
            CV._doc2docx_soffice = lambda src, out_dir: None
            with tempfile.TemporaryDirectory() as td:
                p = self._write(td, 'legacy.doc', OLE2_MAGIC + b'junk-data')
                text, err = CV.doc2md(p)
                self.assertEqual(text, '')
                self.assertIn('doc-no-engine', err)
        finally:
            CV._doc2docx_word_com, CV._doc2docx_soffice = orig_word, orig_soffice

    @staticmethod
    def _fake_converter(content):
        def convert(src, out_dir):
            from docx import Document

            out = os.path.join(out_dir, os.path.splitext(os.path.basename(src))[0] + '.docx')
            d = Document()
            d.add_paragraph(content)
            d.save(out)
            return out
        return convert

    def test_convert_verbose_doc_via_word_com(self):
        orig_word, orig_soffice = CV._doc2docx_word_com, CV._doc2docx_soffice
        try:
            CV._doc2docx_word_com = self._fake_converter('Word 正文段落')
            CV._doc2docx_soffice = lambda src, out_dir: None
            with tempfile.TemporaryDirectory() as td:
                p = self._write(td, 'legacy.doc', OLE2_MAGIC + b'junk-data')
                text, engine, err = CV.convert_verbose(p)
                self.assertEqual(engine, 'doc', err)
                self.assertIn('Word 正文段落', text)
        finally:
            CV._doc2docx_word_com, CV._doc2docx_soffice = orig_word, orig_soffice

    def test_convert_verbose_doc_soffice_fallback(self):
        orig_word, orig_soffice = CV._doc2docx_word_com, CV._doc2docx_soffice
        try:
            CV._doc2docx_word_com = lambda src, out_dir: None
            CV._doc2docx_soffice = self._fake_converter('Libre 正文段落')
            with tempfile.TemporaryDirectory() as td:
                p = self._write(td, 'legacy.doc', OLE2_MAGIC + b'junk-data')
                text, engine, err = CV.convert_verbose(p)
                self.assertEqual(engine, 'doc', err)
                self.assertIn('Libre 正文段落', text)
        finally:
            CV._doc2docx_word_com, CV._doc2docx_soffice = orig_word, orig_soffice


class _FakeComError(Exception):
    """模拟 pywintypes.com_error：hresult 位于 args[0]。"""

    def __init__(self, hresult, message):
        super().__init__(hresult, message)


RPC_E_CALL_REJECTED = -2147418111        # 0x80010001 调用被呼叫方拒绝
RPC_E_SERVERCALL_RETRYLATER = -2147417846  # 0x8001010A 服务器忙
CO_E_SERVER_EXEC_FAILURE = -2146959355   # 0x80080005 服务器运行失败


class TestWordComTransientRetry(unittest.TestCase):
    """V2.3.8 修复 #11：Word COM 冷启动期 RPC_E_CALL_REJECTED 有界重试。"""

    def test_retry_recovers_from_transient_rejection(self):
        state = {'n': 0}

        def flaky():
            state['n'] += 1
            if state['n'] < 3:
                raise _FakeComError(RPC_E_CALL_REJECTED, '调用被呼叫方拒绝')
            return 'ok'

        self.assertEqual(CV._word_com_retry(flaky, attempts=4, delays=(0, 0, 0)), 'ok')
        self.assertEqual(state['n'], 3)

    def test_retry_recovers_from_server_exec_failure(self):
        state = {'n': 0}

        def flaky():
            state['n'] += 1
            if state['n'] < 2:
                raise _FakeComError(CO_E_SERVER_EXEC_FAILURE, '服务器运行失败')
            return 'ok'

        self.assertEqual(CV._word_com_retry(flaky, attempts=4, delays=(0, 0, 0)), 'ok')
        self.assertEqual(state['n'], 2)

    def test_retry_reraises_non_retryable_error(self):
        calls = []

        def broken():
            calls.append(1)
            raise _FakeComError(-2147352567, '发生意外')

        with self.assertRaises(_FakeComError):
            CV._word_com_retry(broken, attempts=4, delays=(0, 0, 0))
        self.assertEqual(len(calls), 1)

    def test_retry_gives_up_after_attempts(self):
        calls = []

        def always_busy():
            calls.append(1)
            raise _FakeComError(RPC_E_SERVERCALL_RETRYLATER, '服务器忙')

        with self.assertRaises(_FakeComError):
            CV._word_com_retry(always_busy, attempts=3, delays=(0, 0))
        self.assertEqual(len(calls), 3)

    def test_doc2docx_word_com_survives_cold_start_rejection(self):
        import types
        from unittest import mock

        module_keys = ('pythoncom', 'win32com', 'win32com.client')
        real_modules = {k: sys.modules.get(k) for k in module_keys}

        class _FakeDoc:
            def SaveAs2(self, out, FileFormat=0):
                with open(out, 'wb') as f:
                    f.write(b'fake-docx')

            def Close(self, flag):
                pass

        class _FakeDocs:
            def Open(self, path, ReadOnly=False):
                return _FakeDoc()

        class _FakeWord:
            def __init__(self):
                self.Visible = None
                self.DisplayAlerts = None
                self.Documents = _FakeDocs()

            def Quit(self):
                pass

        class _FakeClient:
            tries = 0

            def DispatchEx(self, prog_id):
                _FakeClient.tries += 1
                if _FakeClient.tries == 1:
                    raise _FakeComError(RPC_E_CALL_REJECTED, '调用被呼叫方拒绝')
                return _FakeWord()

        fake_pythoncom = types.ModuleType('pythoncom')
        fake_pythoncom.CoInitialize = lambda: None
        fake_pythoncom.CoUninitialize = lambda: None
        fake_win32com = types.ModuleType('win32com')
        fake_client_mod = types.ModuleType('win32com.client')
        fake_client_mod.DispatchEx = _FakeClient().DispatchEx
        fake_win32com.client = fake_client_mod

        sys.modules['pythoncom'] = fake_pythoncom
        sys.modules['win32com'] = fake_win32com
        sys.modules['win32com.client'] = fake_client_mod
        try:
            with mock.patch.object(CV.time, 'sleep', lambda *_: None):
                with tempfile.TemporaryDirectory() as td:
                    src = os.path.join(td, 'legacy.doc')
                    with open(src, 'wb') as f:
                        f.write(OLE2_MAGIC + b'junk-data')
                    out = CV._doc2docx_word_com(src, td)
                    self.assertIsNotNone(out)
                    self.assertTrue(os.path.isfile(out))
                    self.assertEqual(_FakeClient.tries, 2)
        finally:
            for k, v in real_modules.items():
                if v is None:
                    sys.modules.pop(k, None)
                else:
                    sys.modules[k] = v


class TestFormulaHeuristic(unittest.TestCase):
    def test_datetime_not_formula(self):
        self.assertFalse(CV._looks_like_formula('2026-08-31 14:30'))

    def test_phone_number_not_formula(self):
        self.assertFalse(CV._looks_like_formula('010-1234-5678'))

    def test_plain_sentence_not_formula(self):
        self.assertFalse(CV._looks_like_formula('本系统采用 B/S 架构, 支持 Windows 10/11.'))

    def test_real_formula_detected(self):
        self.assertTrue(CV._looks_like_formula('E = mc^2'))
        self.assertTrue(CV._looks_like_formula('x^2 + y^2 = z^2'))
        self.assertTrue(CV._looks_like_formula('α + β = γ'))

    def test_arithmetic_expression_detected(self):
        self.assertTrue(CV._looks_like_formula('3.14 * r^2'))


class TestPdfCrossPageTableMerge(unittest.TestCase):
    def test_merge_repeated_header(self):
        md = (
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '\n'
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| b | 2 |\n'
        )
        out = CV._merge_split_tables(md)
        self.assertEqual(
            out,
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '| b | 2 |\n',
        )

    def test_merge_continuation_fragment(self):
        # 页 N+1 的片段没有表头：首行被 _data_to_md 误升格为表头并多出一条分隔行
        md = (
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '\n'
            '| c | 3 |\n'
            '| --- | --- |\n'
            '| d | 4 |\n'
        )
        out = CV._merge_split_tables(md)
        self.assertEqual(
            out,
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '| c | 3 |\n'
            '| d | 4 |\n',
        )

    def test_no_merge_different_ncol(self):
        md = (
            '| a | b |\n'
            '| --- | --- |\n'
            '| 1 | 2 |\n'
            '\n'
            '| x | y | z |\n'
            '| --- | --- | --- |\n'
            '| 4 | 5 | 6 |\n'
        )
        self.assertEqual(CV._merge_split_tables(md), md)

    def test_no_merge_text_between_tables(self):
        md = (
            '| a | b |\n'
            '| --- | --- |\n'
            '| 1 | 2 |\n'
            '\n'
            '中间有说明文字\n'
            '\n'
            '| a | b |\n'
            '| --- | --- |\n'
            '| 3 | 4 |\n'
        )
        self.assertEqual(CV._merge_split_tables(md), md)

    def test_merge_three_table_chain(self):
        md = (
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '\n'
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| b | 2 |\n'
            '\n'
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| c | 3 |\n'
        )
        out = CV._merge_split_tables(md)
        self.assertEqual(
            out,
            '| 名 | 值 |\n'
            '| --- | --- |\n'
            '| a | 1 |\n'
            '| b | 2 |\n'
            '| c | 3 |\n',
        )

    def test_fence_protects_pipe_rows(self):
        # 代码围栏打断了相邻性：围栏内的管道行不参与合并
        md = (
            '| a | b |\n'
            '| --- | --- |\n'
            '| 1 | 2 |\n'
            '\n'
            '```\n'
            '| a | b |\n'
            '| --- | --- |\n'
            '| 9 | 9 |\n'
            '```\n'
            '\n'
            '| a | b |\n'
            '| --- | --- |\n'
            '| 3 | 4 |\n'
        )
        self.assertEqual(CV._merge_split_tables(md), md)

    def test_pdf2md_cross_page_table_merge(self):
        import fitz
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'cross.pdf')
            doc = fitz.open()
            for values in (['v1', 'v2'], ['v3', 'v4']):
                page = doc.new_page()
                x0, y0, x1, y1 = 72, 72, 280, 120
                rows, cols = 2, 2
                cw = (x1 - x0) / cols
                ch = (y1 - y0) / rows
                for r in range(rows + 1):
                    page.draw_line((x0, y0 + r * ch), (x1, y0 + r * ch), color=(0, 0, 0), width=0.6)
                for c in range(cols + 1):
                    page.draw_line((x0 + c * cw, y0), (x0 + c * cw, y1), color=(0, 0, 0), width=0.6)
                page.insert_text((x0 + 4, y0 + 14), 'H1', fontsize=9)
                page.insert_text((x0 + cw + 4, y0 + 14), 'H2', fontsize=9)
                page.insert_text((x0 + 4, y0 + ch + 14), values[0], fontsize=9)
                page.insert_text((x0 + cw + 4, y0 + ch + 14), values[1], fontsize=9)
            doc.save(p)
            doc.close()

            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'pdf', err)
            rows = _pipe_rows(text)
            self.assertEqual(rows, [['H1', 'H2'], ['v1', 'v2'], ['v3', 'v4']])
            seps = []
            for line in text.splitlines():
                s = line.strip()
                if s.startswith('|') and s.endswith('|'):
                    cells = [c.strip() for c in s[1:-1].split('|')]
                    if cells and all(c == '---' for c in cells):
                        seps.append(line)
            self.assertEqual(len(seps), 1)


class TestPdfBorderlessTable(unittest.TestCase):
    def test_accept_normal_table(self):
        rows = [['Item', 'Qty'], ['Apple', '3'], ['Banana', '5']]
        self.assertTrue(CV._accept_text_table(rows))

    def test_accept_ignores_none_and_blank_cells(self):
        rows = [['Item', 'Qty'], ['Apple', '3'], [None, ''], ['Banana', '5']]
        self.assertTrue(CV._accept_text_table(rows))

    def test_reject_empty(self):
        self.assertFalse(CV._accept_text_table([]))

    def test_reject_single_column(self):
        rows = [['line one'], ['line two'], ['line three']]
        self.assertFalse(CV._accept_text_table(rows))

    def test_reject_single_row(self):
        self.assertFalse(CV._accept_text_table([['a', 'b']]))

    def test_reject_prose_merged_into_one_cell(self):
        self.assertFalse(CV._accept_text_table([['a long prose paragraph merged into one cell']]))

    def test_reject_only_one_multi_cell_row(self):
        rows = [['h1', 'h2'], ['justonecell']]
        self.assertFalse(CV._accept_text_table(rows))

    def test_pdf2md_borderless_table_detected(self):
        import fitz
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'borderless.pdf')
            doc = fitz.open()
            page = doc.new_page()
            y = 100.0
            for name, qty in (('Item', 'Qty'), ('Apple', '3'), ('Banana', '5'), ('Cherry', '12')):
                page.insert_text((80, y), name, fontsize=11)
                page.insert_text((300, y), qty, fontsize=11)
                y += 26
            doc.save(p)
            doc.close()

            text, engine, err = CV.convert_verbose(p)
            self.assertEqual(engine, 'pdf', err)
            rows = _pipe_rows(text)
            self.assertTrue(rows, '无边框表格应还原为管道表，实际输出：\n' + text)
            flat = [c for row in rows for c in row]
            for token in ('Item', 'Qty', 'Apple', '3', 'Banana', '5'):
                self.assertIn(token, flat)


class TestOcrRoutingAndLimits(unittest.TestCase):
    """修复 #8：ocr_any 路由白名单、PDF 截断提示、引擎错误码统一。"""

    def setUp(self):
        from src.readmd_modules import ocr
        self.ocr = ocr
        ocr._engine_cache.clear()

    def test_ocr_any_routes_pdf_and_image(self):
        """回归保护：.pdf → ocr_pdf_to_md，图片后缀 → ocr_image_to_md。"""
        with patch.object(self.ocr, 'ocr_pdf_to_md', return_value='PDF-MD') as mp:
            self.assertEqual(self.ocr.ocr_any('x.pdf'), 'PDF-MD')
            mp.assert_called_once_with('x.pdf')
        with patch.object(self.ocr, 'ocr_image_to_md', return_value='IMG-MD') as mi:
            self.assertEqual(self.ocr.ocr_any('x.png'), 'IMG-MD')
            mi.assert_called_once_with('x.png')

    def test_ocr_any_rejects_non_image_type(self):
        """非图片/PDF 后缀应被 ocr_any 拒绝并返回稳定错误码。"""
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'note.txt')
            with open(p, 'w', encoding='utf-8') as f:
                f.write('hello')
            with patch.object(self.ocr, '_ocr_bytes', return_value=''):
                with self.assertRaises(ValueError) as ctx:
                    self.ocr.ocr_any(p)
            self.assertIn('ocr-unsupported-type', str(ctx.exception))

    def test_ocr_pdf_truncation_notice(self):
        """超过 max_pages 时输出应包含截断提示行。"""
        import fitz
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'multi.pdf')
            doc = fitz.open()
            for i in range(3):
                page = doc.new_page()
                page.insert_text((72, 100), 'page %d content' % (i + 1), fontsize=11)
            doc.save(p)
            doc.close()

            md = self.ocr.ocr_pdf_to_md(p, max_pages=2)
            self.assertIn('## 第 1 页', md)
            self.assertIn('## 第 2 页', md)
            self.assertNotIn('第 3 页', md)
            self.assertIn('仅处理前 2 页', md)

    def test_ocr_pdf_no_notice_within_limit(self):
        """未超限时不输出截断提示。"""
        import fitz
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'one.pdf')
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 100), 'page 1 content', fontsize=11)
            doc.save(p)
            doc.close()

            md = self.ocr.ocr_pdf_to_md(p, max_pages=200)
            self.assertIn('## 第 1 页', md)
            self.assertNotIn('仅处理前', md)

    def test_engine_unavailable_error_code(self):
        """无可用引擎时 _ocr_bytes 应报 ocr-no-engine 错误码。"""
        with patch.object(self.ocr, '_pick_engine', return_value=None):
            with self.assertRaises(RuntimeError) as ctx:
                self.ocr._ocr_bytes(b'x')
        self.assertIn('ocr-no-engine', str(ctx.exception))

    def test_tesseract_missing_error_code(self):
        """Tesseract 缺失时应报 ocr-no-engine 错误码。"""
        with patch('subprocess.run', side_effect=FileNotFoundError):
            with self.assertRaises(RuntimeError) as ctx:
                self.ocr._tesseract_ocr_bytes(b'x')
        self.assertIn('ocr-no-engine', str(ctx.exception))


class TestHtmlConversionRouting(unittest.TestCase):
    """修复 #9：.html/.htm 不再被当源码包 fence，而是经 MarkItDown 转为真正的 Markdown。"""

    def test_html_exts_not_treated_as_code(self):
        """.html/.htm 不应出现在 EXT_TO_LANG（否则会被 code2md 整页包成代码块）。"""
        self.assertNotIn('.html', CV.EXT_TO_LANG)
        self.assertNotIn('.htm', CV.EXT_TO_LANG)

    def test_html_converts_via_markitdown(self):
        """含标题/表格/代码块的 HTML 应转为 Markdown 结构，而非整页源码。"""
        html = (
            '<html><head><title>t</title></head><body>'
            '<h1>报告标题</h1>'
            '<table><tr><th>项目</th><th>数值</th></tr>'
            '<tr><td>甲</td><td>1</td></tr></table>'
            '<pre><code>print(1)</code></pre>'
            '</body></html>'
        )
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'sample.html')
            with open(p, 'w', encoding='utf-8') as f:
                f.write(html)
            text, engine, err = CV.convert_verbose(p)
        self.assertIsNone(err)
        self.assertEqual(engine, 'markitdown')
        self.assertIn('# 报告标题', text)
        self.assertIn('| 项目', text)
        self.assertIn('| ---', text)
        self.assertIn('```', text)
        self.assertNotIn('<h1>', text)

    def test_html_conversion_fallback_without_markitdown(self):
        """MarkItDown 不可用时回退 code_fallback（源码 fence），而非报错。"""
        html = '<html><body><h1>hello</h1></body></html>'
        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, 'sample.htm')
            with open(p, 'w', encoding='utf-8') as f:
                f.write(html)
            with patch.object(CV, '_markitdown_convert', side_effect=ImportError('MarkItDown 未安装')):
                text, engine, err = CV.convert_verbose(p)
        self.assertIsNone(err)
        self.assertEqual(engine, 'code_fallback')
        self.assertIn('```', text)


if __name__ == '__main__':
    unittest.main()
